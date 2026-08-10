#!/usr/bin/env python3
"""Step 1 (AMFM) — leaf → STLA id → CFTS section, by bracket mapping.

The 037's Categorization column is `Functional / NA` on all 102 leaves, so the
RD grouping axis carries no information (framework Part III, §4.1.2 degenerate
case). The axis that does carry it: every 037 Requirement Title ends with its
source STLA id, and every CFTS024 heading carries an `{id}` anchor that is
strictly increasing in document order. A leaf therefore lands in the section
whose anchor is the largest one not exceeding it — a bracket lookup, no
judgement.

Two traps this exists to stop being hand-work:

1. **Full-width brackets.** `SWE-RA-RAD-087` and `-097` tag their id as
   `（4942534）` (U+FF08/U+FF09), not `(4942534)`. An ASCII-only pattern loses
   exactly those two leaves and reports them as untagged.
2. **Out-of-range ids are not misses.** 17 leaves carry ids far above the
   CFTS024 anchor range because they belong to other CFTS documents. Bracket
   lookup would silently pin every one of them to the document's last section.
   Out-of-range is detected and routed by the ruled allocation instead
   (R7-Q3), and the two must agree on the same set of leaves or this fails.

Outputs (`data/`):
  stla_to_cfts.json  leaf -> {stla_id, doc, section, section_title, anchor}
  stla_to_cfts.tsv   the same, greppable, for review

Usage:
    python AMFMHMI/scripts/build_stla_map.py --feature-dir AMFMHMI
    python AMFMHMI/scripts/build_stla_map.py --feature-dir AMFMHMI --check-batches
"""
from __future__ import annotations

import argparse
import bisect
import difflib
import json
import re
import sys
from pathlib import Path

import docx
import openpyxl

sys.path.insert(0, str(Path(__file__).resolve().parent.parent.parent / "scripts"))
from feature_config import load_feature_config, resolve_path  # noqa: E402

# Both bracket families, because the 037 mixes them. Curly braces appear in
# the CFTS headings; the 037 uses round or full-width round.
STLA_ID_RE = re.compile(r"[({（]\s*(\d{7})\s*[)}）]")
HEADING_RE = re.compile(
    r"^(?P<num>\d+(?:\.\d+)*)\s+(?P<title>.+?)\s*[{(（]\s*(?P<anchor>\d{7})\s*[)}）]\s*$")
# Body requirements carry their own id as a leading tag on the metadata line;
# the requirement sentence is the paragraph(s) that follow it.
PARA_ANCHOR_RE = re.compile(r"^(\d{7})\s*[:：]\s*(.*)$")
A03_SHEET = "Analysis Report"


class BuildError(RuntimeError):
    pass


def norm(v) -> str:
    return re.sub(r"\s+", " ", str(v or "")).strip().lower()


# ------------------------------------------------------------------- inputs

def load_leaves(a03_path: Path) -> list[dict]:
    """037 leaves in document order: req_id, title, description, stla_id."""
    wb = openpyxl.load_workbook(a03_path, read_only=True)
    rows = list(wb[A03_SHEET].iter_rows(values_only=True))
    wb.close()
    hdr = next((i for i, r in enumerate(rows)
                if any("requirement description" in norm(v) for v in r)), None)
    if hdr is None:
        raise BuildError(f"{a03_path.name}: no header row")
    header = [norm(v) for v in rows[hdr]]

    def col(*need, forbid=()):
        hits = [i for i, v in enumerate(header)
                if all(t in v for t in need) and not any(t in v for t in forbid)]
        if len(hits) != 1:
            raise BuildError(f"{a03_path.name}: {need} matched {len(hits)} columns")
        return hits[0]

    ti, di = col("requirement", "title"), col("requirement", "description")
    ci = col("categorization", forbid=("sub",))
    si = col("source", "id")

    leaves, untagged, ambiguous = [], [], []
    for r in rows[hdr + 1:]:
        if not r[0] or not str(r[ci] or "").strip().lower().startswith("functional"):
            continue
        rid, title = str(r[0]).strip(), str(r[ti] or "")
        ids = STLA_ID_RE.findall(title)
        if not ids:
            untagged.append(rid)
            continue
        # The id is a trailing tag; a title quoting another id earlier would
        # be a genuine ambiguity, so distinct values are reported, not picked.
        if len(set(ids)) > 1:
            ambiguous.append((rid, sorted(set(ids))))
            continue
        leaves.append({"req_id": rid, "stla_id": int(ids[-1]),
                       "declared_stla_id": int(ids[-1]),
                       "title": title.strip(),
                       "description": str(r[di] or "").strip(),
                       "source_components": str(r[si] or "").strip()})
    if untagged or ambiguous:
        raise BuildError(
            "cannot read an STLA id from every leaf — this map must be total, "
            f"not best-effort. untagged={untagged} ambiguous={ambiguous}")
    return leaves


def load_paragraph_anchors(docx_path: Path) -> dict[int, dict]:
    """{stla id: {metadata, text}} for requirements written in the body.

    Every CFTS requirement paragraph is introduced by a metadata line tagged
    with its own id (`4872375: [Artifact Type:…] [State:…]`); the requirement
    sentence is what follows, up to the next tagged line or heading. This is a
    strictly finer resolution than the section bracket, and all 85 in-range
    leaves reach it — so a leaf's context can be its own clause rather than
    its chapter.
    """
    paras = docx.Document(docx_path).paragraphs
    tagged = [(i, PARA_ANCHOR_RE.match(p.text.strip()))
              for i, p in enumerate(paras)]
    tagged = [(i, m) for i, m in tagged if m]
    out = {}
    for k, (i, m) in enumerate(tagged):
        end = tagged[k + 1][0] if k + 1 < len(tagged) else len(paras)
        body = []
        for p in paras[i + 1:end]:
            if p.style.name.startswith("Heading"):
                break
            if p.text.strip():
                body.append(p.text.strip())
        out[int(m.group(1))] = {"metadata": m.group(2).strip(),
                                "text": "\n".join(body)}
    return out


def load_sections(docx_path: Path) -> tuple[list[dict], list[str]]:
    """CFTS headings in document order, with their anchors."""
    doc = docx.Document(docx_path)
    sections, unparsed = [], []
    for p in doc.paragraphs:
        if not p.style.name.startswith("Heading"):
            continue
        text = p.text.strip()
        if not text:
            continue
        m = HEADING_RE.match(text)
        if m:
            sections.append({"section": m.group("num"),
                             "title": m.group("title").strip(),
                             "anchor": int(m.group("anchor"))})
        else:
            unparsed.append(text)
    if not sections:
        raise BuildError(f"{docx_path.name}: no anchored headings found")
    anchors = [s["anchor"] for s in sections]
    breaks = [(sections[i]["section"], anchors[i], sections[i + 1]["section"],
               anchors[i + 1])
              for i in range(len(anchors) - 1) if anchors[i] >= anchors[i + 1]]
    if breaks:
        raise BuildError(
            "CFTS heading anchors are not strictly increasing in document "
            "order, so a bracket lookup is not defined. Breaks: "
            f"{breaks[:5]}")
    return sections, unparsed


# ------------------------------------------------------------------ mapping

def parse_leaf_selector(spec: str, prefix: str) -> list[str]:
    """`087, 089-096` -> full req_ids. Ranges are inclusive and numeric."""
    out = []
    for part in str(spec).split(","):
        part = part.strip()
        if not part:
            continue
        if "-" in part:
            a, b = (int(x) for x in part.split("-", 1))
            out.extend(f"{prefix}{n:03d}" for n in range(a, b + 1))
        else:
            out.append(f"{prefix}{int(part):03d}")
    return out


def bracket(sections: list[dict], stla_id: int) -> dict | None:
    """The section whose anchor is the largest not exceeding stla_id."""
    anchors = [s["anchor"] for s in sections]
    i = bisect.bisect_right(anchors, stla_id) - 1
    return sections[i] if i >= 0 else None


def build(leaves, sections, external: dict[str, list[str]], home_doc: str,
          paragraphs: dict[int, dict] | None = None):
    lo, hi = sections[0]["anchor"], sections[-1]["anchor"]
    declared_external = {rid: doc for doc, ids in external.items() for rid in ids}

    mapping, out_of_range = {}, []
    for leaf in leaves:
        sid = leaf["stla_id"]
        inside = lo <= sid <= hi
        if not inside:
            out_of_range.append(leaf["req_id"])
        entry = {"stla_id": sid, "title": leaf["title"],
                 "description": leaf["description"],
                 "source_components": leaf["source_components"]}
        if inside:
            sec = bracket(sections, sid)
            para = (paragraphs or {}).get(sid)
            entry |= {"doc": home_doc, "section": sec["section"],
                      "section_title": sec["title"], "anchor": sec["anchor"],
                      "spec_paragraph": para["text"] if para else None,
                      "spec_paragraph_metadata": para["metadata"] if para else None,
                      "resolution": "paragraph" if para else "bracket"}
        else:
            doc = declared_external.get(leaf["req_id"])
            entry |= {"doc": doc, "section": None, "section_title": None,
                      "anchor": None,
                      "resolution": "external-allocation" if doc else "UNRESOLVED"}
        mapping[leaf["req_id"]] = entry

    # The mechanical test (is the id inside this document's anchor range?) and
    # the ruled allocation (which other document it belongs to) are derived
    # independently. They must select the same leaves; a disagreement means one
    # of them moved, and guessing which would silently mis-cite a spec.
    declared = set(declared_external)
    found = set(out_of_range)
    if declared != found:
        raise BuildError(
            "the ruled external allocation and the measured out-of-range set "
            f"disagree.\n  declared external, but inside {home_doc}: "
            f"{sorted(declared - found)}\n  outside {home_doc}, but not "
            f"allocated to any document: {sorted(found - declared)}")
    return mapping


def apply_overrides(leaves: list[dict], overrides: dict,
                    paragraphs: dict[int, dict]) -> list[dict]:
    """Apply ruled declared-id corrections, refusing ones that do not hold up.

    An override rewrites which clause a requirement traces to, so it is checked
    against the same evidence that justified it: the corrected clause must
    match the leaf's own text better than the declared one does. A ruling that
    stops being true of the files — the 037 reissued, the CFTS renumbered —
    then fails here instead of silently re-pointing a citation.
    """
    applied = []
    by_id = {l["req_id"]: l for l in leaves}
    for rid, spec in (overrides or {}).items():
        leaf = by_id.get(rid)
        if leaf is None:
            raise BuildError(f"stla_id_overrides names {rid}, which is not a "
                             "037 leaf")
        declared, corrected = int(spec["declared"]), int(spec["corrected"])
        if leaf["declared_stla_id"] != declared:
            raise BuildError(
                f"{rid}: override expects declared id {declared}, but the 037 "
                f"now says {leaf['declared_stla_id']} — the ruling "
                f"({spec.get('ruling', '?')}) predates this file")
        if corrected not in paragraphs:
            raise BuildError(f"{rid}: override target {corrected} is not a "
                             "clause in the primary document")
        was = difflib.SequenceMatcher(
            None, _fold(leaf["title"]),
            _fold(paragraphs.get(declared, {}).get("text", ""))).ratio()
        now = difflib.SequenceMatcher(
            None, _fold(leaf["title"]),
            _fold(paragraphs[corrected]["text"])).ratio()
        if now <= was:
            raise BuildError(
                f"{rid}: override to {corrected} agrees {now:.3f} with the "
                f"leaf, no better than the declared {declared} at {was:.3f} — "
                "the evidence for this ruling no longer holds")
        leaf["stla_id"] = corrected
        applied.append({"req_id": rid, "declared": declared,
                        "corrected": corrected, "ruling": spec.get("ruling"),
                        "agreement_declared": round(was, 3),
                        "agreement_corrected": round(now, 3)})
    return applied


def _fold(s: str) -> str:
    s = re.sub(r"[({（]\s*\d{7}\s*[)}）]", " ", str(s or ""))
    return re.sub(r"\s+", " ", re.sub(r"[^a-z0-9$ ]", " ", s.lower())).strip()


def verify_ids(mapping: dict, paragraphs: dict[int, dict],
               floor: float = 0.55, margin: float = 0.15) -> list[dict]:
    """Does each leaf's declared id point at the clause the leaf describes?

    The id is a hand-typed tail on the 037 title, so it is copyable — and
    A-AM08 already records four suspected duplicates. A wrong id is invisible
    to the bracket map (it resolves perfectly, to the wrong clause) and ships
    as a wrong spec_reference, which is a traceability defect rather than a
    wording one.

    Reported when the declared paragraph is a poor match AND some other
    paragraph is clearly better. Both conditions matter: a leaf that
    paraphrases heavily scores low against everything and is not evidence of
    a wrong id.

    A token-overlap prefilter keeps this to a few seconds over ~1400
    paragraphs; difflib then ranks the survivors.
    """
    index = {pid: (_fold(p["text"]), set(_fold(p["text"]).split()))
             for pid, p in paragraphs.items()}
    findings = []
    for rid, e in mapping.items():
        if not e.get("spec_paragraph"):
            continue
        want = _fold(e["title"])
        want_set = set(want.split())
        if not want_set:
            continue
        declared = difflib.SequenceMatcher(
            None, want, index[e["stla_id"]][0]).ratio()
        if declared >= floor:
            continue
        prefiltered = sorted(
            index, key=lambda pid: -len(want_set & index[pid][1]))[:20]
        best_pid, best = None, 0.0
        for pid in prefiltered:
            r = difflib.SequenceMatcher(None, want, index[pid][0]).ratio()
            if r > best:
                best_pid, best = pid, r
        if best_pid != e["stla_id"] and best - declared >= margin:
            findings.append({
                "req_id": rid, "declared_id": e["stla_id"],
                "declared_agreement": round(declared, 3),
                "better_id": best_pid, "better_agreement": round(best, 3),
            })
    return findings


def check_batches(mapping: dict, batches_md: Path) -> list[str]:
    """Every leaf in a batch must bracket into one of that batch's sections.

    The Test Set table was derived by hand at Phase 3. This turns it into a
    checked claim: the batch doc says which CFTS sections a set covers, the
    bracket map says which section each leaf actually lands in.
    """
    row = re.compile(r"^\|\s*([^|]+?)\s*\|\s*([^|]*?)\s*\|\s*(\d+)\s*\|"
                     r"\s*([^|]+?)\s*\|")
    problems, seen = [], set()
    for line in batches_md.read_text(encoding="utf-8").splitlines():
        m = row.match(line)
        if not m or m.group(3) == "n":
            continue
        name, secs, count, ids = m.groups()
        req_ids = parse_leaf_selector(ids, "SWE-RA-RAD-")
        if len(req_ids) != int(count):
            problems.append(f"{name}: table says {count} leaves, lists {len(req_ids)}")
        declared = {s.strip() for s in re.split(r"[,;]", secs)
                    if re.match(r"^\s*\d+(\.\d+)*\s*$", s)}
        for rid in req_ids:
            seen.add(rid)
            entry = mapping.get(rid)
            if entry is None:
                problems.append(f"{name}: {rid} is not a 037 leaf")
            elif entry["section"] and declared and entry["section"] not in declared:
                problems.append(
                    f"{name}: {rid} (id {entry['stla_id']}) brackets into "
                    f"{entry['section']} '{entry['section_title']}', which the "
                    f"table does not list ({sorted(declared)})")
    missing = sorted(set(mapping) - seen)
    if missing:
        problems.append(f"{len(missing)} leaves are in no batch: {missing[:10]}")
    return problems


# -------------------------------------------------------------------- main

def run(args) -> int:
    cfg = load_feature_config(args.feature_dir)
    root = Path(args.feature_dir)
    a03 = resolve_path(cfg, "a03_report")
    cfts = resolve_path(cfg, "cfts_doc")

    spec_cfg = cfg.get("spec_docs") or {}
    home_doc = spec_cfg.get("primary")
    external = {d: parse_leaf_selector(v, "SWE-RA-RAD-")
                for d, v in (spec_cfg.get("external") or {}).items()}
    if not home_doc:
        raise BuildError("feature.yaml needs spec_docs.primary (e.g. CFTS024)")

    leaves = load_leaves(a03)
    sections, unparsed = load_sections(cfts)
    paragraphs = load_paragraph_anchors(cfts)
    overrides = apply_overrides(leaves, cfg.get("stla_id_overrides"), paragraphs)
    mapping = build(leaves, sections, external, home_doc, paragraphs)
    for o in overrides:
        mapping[o["req_id"]] |= {
            "declared_stla_id": o["declared"], "id_override": o}

    data = root / "data"
    data.mkdir(exist_ok=True)
    (data / "stla_to_cfts.json").write_text(
        json.dumps(mapping, ensure_ascii=False, indent=1), encoding="utf-8")
    lines = ["req_id\tstla_id\tdoc\tsection\tsection_title\tresolution"]
    for rid, e in mapping.items():
        lines.append(f"{rid}\t{e['stla_id']}\t{e['doc']}\t{e['section'] or ''}"
                     f"\t{e['section_title'] or ''}\t{e['resolution']}")
    (data / "stla_to_cfts.tsv").write_text("\n".join(lines) + "\n", encoding="utf-8")

    by_doc: dict[str, int] = {}
    for e in mapping.values():
        by_doc[e["doc"]] = by_doc.get(e["doc"], 0) + 1
    print(f"leaves        : {len(mapping)}")
    print(f"{home_doc} sections: {len(sections)} anchored headings "
          f"({sections[0]['anchor']}–{sections[-1]['anchor']})")
    if unparsed:
        print(f"headings without an anchor (not usable as brackets): "
              f"{len(unparsed)}")
    for doc, n in sorted(by_doc.items(), key=lambda kv: -kv[1]):
        print(f"  {doc}: {n} leaves")
    print(f"distinct {home_doc} sections used: "
          f"{len({e['section'] for e in mapping.values() if e['section']})}")
    by_res: dict[str, int] = {}
    for e in mapping.values():
        by_res[e["resolution"]] = by_res.get(e["resolution"], 0) + 1
    print(f"paragraph anchors in {home_doc}: {len(paragraphs)}")
    print("resolution    : " + ", ".join(f"{k}={v}" for k, v in sorted(by_res.items())))
    coarse = [rid for rid, e in mapping.items() if e["resolution"] == "bracket"]
    if coarse:
        print(f"  section-level only (no paragraph anchor for the id): {coarse}")

    for o in overrides:
        print(f"id override   : {o['req_id']} {o['declared']} -> "
              f"{o['corrected']} ({o['ruling']}); agreement "
              f"{o['agreement_declared']} -> {o['agreement_corrected']}")
    suspect = verify_ids(mapping, paragraphs)
    if suspect:
        print("\nDECLARED STLA ID DOES NOT MATCH THE CLAUSE (A-AM08 class):",
              file=sys.stderr)
        for f in suspect:
            print(f"  - {f['req_id']}: declared {f['declared_id']} "
                  f"(agreement {f['declared_agreement']}), but {f['better_id']} "
                  f"matches at {f['better_agreement']}", file=sys.stderr)
        print("  These are reported, never corrected here — R7-Q4 leaves each "
              "pair to Pei.", file=sys.stderr)
        (data / "stla_id_suspects.json").write_text(
            json.dumps(suspect, ensure_ascii=False, indent=1), encoding="utf-8")

    if args.check_batches:
        md = root / args.batches_md
        if not md.exists():
            raise BuildError(f"missing {md}")
        problems = check_batches(mapping, md)
        if problems:
            print("\nBATCH TABLE DISAGREES WITH THE BRACKET MAP:", file=sys.stderr)
            for p in problems:
                print(f"  - {p}", file=sys.stderr)
            return 1
        print(f"batch table   : consistent with the bracket map ({md.name})")
    return 0


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--feature-dir", default=".")
    ap.add_argument("--batches-md", default="docs/batches-amfm.md")
    ap.add_argument("--check-batches", action="store_true")
    args = ap.parse_args()
    try:
        return run(args)
    except BuildError as exc:
        print(f"ABORTED: {exc}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    sys.exit(main())
