#!/usr/bin/env python3
"""R-DM8 lookup: do CFTS_020 / SYS3 SYSAD carry the four missing values?

R-DM8 requires each gap to be searched in the CFTS body and the SYS3 SYSAD
FIRST; only what is not found there may be registered as DR-DM. This script
reports HITS WITH LOCATION, never a value judgement — reading the numbers out
of a hit and deciding they are the threshold is Phase 2 work.

Location is given as the nearest preceding Heading-styled paragraph, so a hit
can be cited by clause rather than by paragraph index.
"""
import re
from pathlib import Path

import docx

R = Path(__file__).resolve().parents[1] / "inputs"
DOCS = {
    "CFTS_020": R / ("R1LR_Atl-H_26PI1.5 Mar Release-Cabin_CFTS_020 "
                     "ICS and DCSD _20260310-1533.docx"),
    "SYS3_SYSAD": R / ("SYS3_CFTS_020_display_FM-WI-FSM-011-A01_System "
                       "Architectural Design_SYSAD_v1.0.docx"),
}

GAPS = {
    "SWE-DM-003 splash/sleep 時長門檻": r"splash|sleep",
    "SWE-DM-004 thermal warning threshold": r"thermal|temperature|overheat|°\s*c|degc",
    "SWE-DM-005 thermal protection critical/回復": r"critical|protection|shut\s*down|recover",
    "SWE-DM-006 popup priority / timeout": r"priorit|arbitrat|time\s*out|timeout",
}
NUM = re.compile(r"\b\d+(?:\.\d+)?\s*(?:ms|s|sec|second|min|minute|°|deg|c\b)",
                 re.I)


def blocks(path):
    d = docx.Document(path)
    out, head = [], "(前言)"
    for p in d.paragraphs:
        style = str(p.style.name).lower() if p.style is not None else ""
        txt = " ".join(p.text.split())
        if style.startswith("heading") and txt:
            head = txt
        if txt:
            out.append((head, txt))
    for t in d.tables:
        for row in t.rows:
            for c in row.cells:
                txt = " ".join(c.text.split())
                if txt:
                    out.append(("(表格)", txt))
    return out


def main():
    print("# R-DM8 lookup — 四處缺值在 CFTS / SYS3 之查證")
    print("方法：段落層 regex；location = 最近之 Heading 樣式段落")
    print("本腳本只報命中與其位置，不讀出任何數值作為門檻（Phase 2）")
    for name, path in DOCS.items():
        bl = blocks(path)
        print(f"\n## {name} ({len(bl)} 段/格)")
        for gap, pat in GAPS.items():
            rx = re.compile(pat, re.I)
            hits = [(h, t) for h, t in bl if rx.search(t)]
            numeric = [(h, t) for h, t in hits if NUM.search(t)]
            heads = sorted({h for h, _ in hits})
            print(f"\n### {gap}")
            print(f"  命中段落 {len(hits)}；其中含數值+單位者 {len(numeric)}")
            print(f"  涉及章節 {len(heads)} 個：{heads[:12]}"
                  + (" …" if len(heads) > 12 else ""))
            for h, t in numeric[:5]:
                print(f"    [{h}] {t[:150]}")


if __name__ == "__main__":
    main()
