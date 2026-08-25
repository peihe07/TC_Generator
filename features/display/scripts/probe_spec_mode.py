#!/usr/bin/env python3
"""spec_mode extraction-capability probe (handoff 01 step 11, canon §3).

Measures what each of the three source forms can actually yield, so the
spec_mode choice rests on measurement rather than on the scaffold's proposal:

  CFTS_020 docx  — paragraph/table yield, clause-id inventory, and whether an
                   id found in the body is also reachable from a heading
                   (index) — the 全文 id 數 vs 索引數 comparison
  SYS3 SYSAD docx — same, plus image count (a SYSAD is often mostly figures)
  SYS2 xlsx      — Basic Report field structure already measured in
                   recount_sys2.py; here only the outline/section-anchor
                   question is asked
"""
import re
import zipfile
from collections import Counter
from pathlib import Path

import docx
import openpyxl

from tsv_meta import write_meta

FEAT = Path(__file__).resolve().parents[1]
ROOT = FEAT
CFTS = ROOT / "inputs" / ("R1LR_Atl-H_26PI1.5 Mar Release-Cabin_CFTS_020 "
                          "ICS and DCSD _20260310-1533.docx")
SYS3 = ROOT / "inputs" / ("SYS3_CFTS_020_display_FM-WI-FSM-011-A01_System "
                          "Architectural Design_SYSAD_v1.0.docx")
SYS2 = ROOT / "inputs" / ("SYS2_CFTS_020_DISP_TCH_ICS_20260616_All_HW_System"
                          "_Accepted & Released.xlsx")

# Clause-id shapes observed in this corpus. PSCFTS020-1-45-1 is the Melco ID
# form that 037's Excluded sheet carries (R-DM4); the outline form is the
# numbered-heading form a spec_reference would be built from under mode B/D.
ID_PATTERNS = {
    "melco (PSCFTS020-n-n-n)": r"PSCFTS\d{3}-\d+-\d+(?:-\d+)?",
    "outline heading (n.n / n.n.n)": r"(?m)^\s*(\d+(?:\.\d+){1,3})\s+\S",
}


def norm(s):
    return " ".join(str(s or "").split())


def spec_text_layer_rows():
    """The character counts this file yields under each extractor.

    Three numbers for one file, because `recon.py`'s pipeline probe and this
    script measure it differently — 854,333 (pymupdf) vs 907,382
    (python-docx). Neither is wrong; they extract different things from a
    .docx. The registered value is the pipeline's (下放包 10 §2.1), and the
    others are kept here so the difference stays visible instead of living
    in one terminal session's scrollback.
    """
    name = CFTS.name
    rows = []
    try:
        import pymupdf as fitz
    except ImportError:
        try:
            import fitz
        except ImportError:
            fitz = None
    if fitz is not None:
        doc = fitz.open(CFTS)
        rows.append(("spec_pdf", name, "pymupdf",
                     str(sum(len(pg.get_text()) for pg in doc)), "Y",
                     "recon.py survey_spec_text_layer()；跨 feature 之管線探針"))
        doc.close()
    d = docx.Document(CFTS)
    paras = [norm(p.text) for p in d.paragraphs if p.text.strip()]
    cells = [norm(c.text) for t in d.tables for r in t.rows for c in r.cells
             if c.text.strip()]
    rows.append(("spec_pdf", name, "python-docx（段落＋表格格，正規化後）",
                 str(len("\n".join(paras + cells))), "N",
                 "features/display/scripts/probe_spec_mode.py；本 feature 自測"))
    raw = "\n".join([p.text for p in d.paragraphs]
                    + [c.text for t in d.tables for r in t.rows for c in r.cells])
    rows.append(("spec_pdf", name, "python-docx（未正規化、含空段）",
                 str(len(raw)), "N", "同上，另一種計法"))
    return rows


def probe_docx(path, label):
    print(f"\n## {label} — {path.name}")
    d = docx.Document(path)
    paras = [p.text for p in d.paragraphs]
    nonempty = [p for p in paras if p.strip()]
    print(f"paragraphs: {len(paras)} (non-empty {len(nonempty)})")
    print(f"tables: {len(d.tables)}")
    cells = []
    for t in d.tables:
        for row in t.rows:
            for c in row.cells:
                if c.text.strip():
                    cells.append(c.text)
    print(f"non-empty table cells: {len(cells)}")
    with zipfile.ZipFile(path) as z:
        media = [n for n in z.namelist() if n.startswith("word/media/")]
    print(f"embedded media files: {len(media)}")

    body = "\n".join(nonempty + cells)
    print(f"extracted characters: {len(body)}")

    heads = [p for p in d.paragraphs
             if p.style is not None and str(p.style.name).lower()
             .startswith("heading")]
    print(f"paragraphs with a Heading style: {len(heads)}")
    head_text = "\n".join(h.text for h in heads)

    for name, pat in ID_PATTERNS.items():
        full = re.findall(pat, body)
        idx = re.findall(pat, head_text)
        print(f"  id form {name}: 全文 {len(full)} 次 / "
              f"{len(set(full))} 相異 | 標題(索引) {len(set(idx))} 相異")
        if set(full):
            sample = sorted(set(full))[:5]
            print(f"     sample: {sample}")
    return body


def main():
    print("# spec_mode extraction probe")
    print("engine: python-docx (paragraphs + table cells), zipfile for media")

    cfts = probe_docx(CFTS, "CFTS_020 本文 (candidate spec source, mode D)")
    sys3 = probe_docx(SYS3, "SYS3 SYSAD (candidate, traceability/architecture)")

    print("\n## SYS2 Polarion export — outline anchor availability")
    wb = openpyxl.load_workbook(SYS2, read_only=True, data_only=True)
    grid = [list(r) for r in wb["Basic Report"].iter_rows(values_only=True)]
    wb.close()
    head = [" ".join(str(h or "").split()) for h in grid[0]]
    doc_col = [i for i, h in enumerate(head) if "Document ID" in h]
    src_col = [i for i, h in enumerate(head) if "Source Requirement items" in h]
    print(f"columns present: Document ID={bool(doc_col)} "
          f"Source Requirement items={bool(src_col)}")
    for label, cols in (("Document ID", doc_col),
                        ("Source Requirement items", src_col)):
        if not cols:
            continue
        vals = [" ".join(str(r[cols[0]] or "").split()) for r in grid[1:]
                if str(r[0] or "").strip()]
        nonblank = [v for v in vals if v]
        print(f"  {label}: {len(nonblank)}/{len(vals)} non-blank, "
              f"{len(set(nonblank))} distinct")
        for v, n in Counter(nonblank).most_common(5):
            print(f"     {v!r} x{n}")

    cols = ["path_key", "file", "extractor", "chars", "registered", "source"]
    rows = spec_text_layer_rows()

    # R-G23 之同分寸，施於自己的量測：the sidecar records what each
    # extractor yielded last time. If a library update moves a number, say
    # so and print both — do not quietly overwrite the expectation, because
    # a file that silently follows whatever it measures asserts nothing.
    import json as _json
    meta_path = (FEAT / "data" / "spec_text_layer.tsv").with_suffix(
        ".tsv.meta.json")
    expected = {}
    if meta_path.is_file():
        expected = (_json.loads(meta_path.read_text(encoding="utf-8"))
                    .get("expected_chars") or {})
    drift = [(r[2], expected[r[2]], r[3]) for r in rows
             if r[2] in expected and expected[r[2]] != r[3]]
    for extractor, want, got in drift:
        print(f"WARNING: spec text layer drift — {extractor}: "
              f"sidecar records {want}, this run measured {got}. "
              f"NOT updating the expectation (R-G23 分寸：察覺變動，"
              f"不代為採納)。")
    out = FEAT / "data" / "spec_text_layer.tsv"
    with out.open("w", encoding="utf-8") as fh:
        fh.write("\t".join(cols) + "\n")
        for r in rows:
            fh.write("\t".join(r) + "\n")
    write_meta(out, cols, len(rows),
               generated_by="features/display/scripts/probe_spec_mode.py",
               rulings=["R-G19", "R-DM34", "A-DM26"],
               measurement_conditions=(
                   "同一份 .docx 由三種計法量得。三數皆由本腳本現算，"
                   "非人工登記；差異為抽取器不同而非計算錯誤。"),
               # Keep whatever the sidecar already recorded for an
               # extractor; only add keys it has never seen. Rewriting it
               # from this run would adopt the drift the WARNING above just
               # refused to adopt — the message would say one thing and the
               # file would do the other.
               expected_chars={**{r[2]: r[3] for r in rows}, **expected},
               notes=(
                   "登記值取 pymupdf（下放包 10 §2.1）：該數字之用途是判斷"
                   "有無文字層（門檻 500 字元），三數皆遠超門檻，導出之結論"
                   "相同（spec_mode D 成立）；跨 feature 之可比性由管線探針"
                   "提供。另記 A-DM26：欄名為 spec_pdf 而內容為 .docx。"))
    print("\n## spec text layer —— 三種計法（寫入 data/spec_text_layer.tsv）")
    for r in rows:
        print(f"  {r[2]:38s} {r[3]:>8}  registered={r[4]}")

    print("\n## 交叉：SYS2 Melco ID 是否可在 CFTS 本文中定位")
    wb = openpyxl.load_workbook(SYS2, read_only=True, data_only=True)
    grid = [list(r) for r in wb["Basic Report"].iter_rows(values_only=True)]
    wb.close()
    mcol = [i for i, h in enumerate(head) if h == "SYS2 Melco ID"][0]
    toks = set()
    for r in grid[1:]:
        if str(r[0] or "").strip():
            for t in re.split(r"[,\s;]+", str(r[mcol] or "")):
                if t.strip():
                    toks.add(t.strip())
    hit = sum(1 for t in toks if t in cfts)
    print(f"  SYS2 Melco tokens: {len(toks)}; found verbatim in CFTS body: {hit}")
    hit3 = sum(1 for t in toks if t in sys3)
    print(f"  found verbatim in SYS3 body: {hit3}")


if __name__ == "__main__":
    main()
