#!/usr/bin/env python3
"""R-DM22 glossary build (handoff 06 step 2).

An entry is created ONLY where an abbreviation and its expansion appear
side by side in one sentence of one source. Two verbatim shapes are swept:

    Expansion (ABBR)      e.g. "Rear View Camera (RVC)"
    ABBR (Expansion)      e.g. "RVC (Rear View Camera)"

No entry is created from domain knowledge — that would be canon §8.4.1.
Per R-DM22(a) an expansion of fewer than two words is not usable as an
anchor; such pairs are still recorded, marked usable=N.

Sources swept in full: 037 (all sheets, all cells), SYS2 `Basic Report`
(all cells), CFTS_020 and SYS3 SYSAD (paragraphs + table cells).

停止條件 16: if one abbreviation gets two different expansions, this stops
and reports rather than choosing.
"""
import re
import sys
from collections import defaultdict
from pathlib import Path

import docx
import openpyxl

from tsv_meta import write_meta

ROOT = Path(__file__).resolve().parents[3]
FEAT = Path(__file__).resolve().parents[1]
F037 = FEAT / "inputs" / \
    "Display_Management_FM-WI-FSM-037-A03_STLA_Report_SWRA.xlsx"
SYS2 = FEAT / "inputs" / ("SYS2_CFTS_020_DISP_TCH_ICS_20260616_All_HW_System"
                          "_Accepted & Released.xlsx")
CFTS = FEAT / "inputs" / ("R1LR_Atl-H_26PI1.5 Mar Release-Cabin_CFTS_020 "
                          "ICS and DCSD _20260310-1533.docx")
SYS3 = FEAT / "inputs" / ("SYS3_CFTS_020_display_FM-WI-FSM-011-A01_System "
                          "Architectural Design_SYSAD_v1.0.docx")

MUST_CHECK = ["DCSD", "ICS", "HU", "FPDM", "LVDS", "SK", "TGW", "SGW", "ETM",
              "RVC"]

# Bracket forms swept: "... words (ABBR)" and "ABBR (words ...)".
# The bracket alone is not enough — "if a high priority screen (RVC)" has
# the same shape as "Rear View Camera (RVC)". The accept test is that the
# candidate words' INITIALS spell the abbreviation, letter for letter:
#
#     Rear View Camera        -> R,V,C  == RVC   accepted
#     high priority screen    -> h,p,s  != RVC   rejected
#
# That is a verbatim test on the source's own characters, not a similarity
# score. FILLER is the one concession — lower-case joining words that a
# writer does not put into an abbreviation ("Body Control Module (BCM)" vs
# "Lost Communication with Radio"); they may be skipped, and whether any
# were is recorded per entry.
PAREN_AFTER = re.compile(r"([A-Za-z][A-Za-z\-/'’]*(?:\s+[A-Za-z][A-Za-z\-/'’]*){0,7})"
                         r"\s*\(\s*([A-Z][A-Za-z0-9]{1,5})\s*\)")
PAREN_BEFORE = re.compile(r"\b([A-Z][A-Za-z0-9]{1,5})\s*\(\s*"
                          r"([A-Za-z][A-Za-z\-/'’]*(?:\s+[A-Za-z][A-Za-z\-/'’]*){0,7})"
                          r"\s*\)")
FILLER = {"of", "and", "the", "for", "with", "to", "a", "an", "in", "on"}


def initials_match(words, abbrev):
    """Do these words' initials spell `abbrev`? Returns (ok, used_filler).

    Tried twice: first with every word counted, then allowing FILLER words
    to be skipped. The caller records which pass succeeded.
    """
    for skip_filler in (False, True):
        ws = [w for w in words
              if not (skip_filler and w.lower() in FILLER)]
        if len(ws) != len(abbrev):
            continue
        if "".join(w[0] for w in ws).upper() == abbrev.upper():
            return True, skip_filler
    return False, False


def longest_expansion(text_words, abbrev):
    """The trailing run of words whose initials spell abbrev, or None."""
    n = len(abbrev)
    for take in range(n, min(len(text_words), n + len(FILLER)) + 1):
        cand = text_words[-take:]
        ok, filler = initials_match(cand, abbrev)
        if ok:
            return " ".join(cand), filler
    return None, False


def norm(s):
    return " ".join(str(s or "").split())


def cells_xlsx(path, sheets=None):
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    for ws in wb.worksheets:
        if sheets and ws.title not in sheets:
            continue
        for i, row in enumerate(ws.iter_rows(values_only=True), 1):
            for j, v in enumerate(row, 1):
                if v is not None and str(v).strip():
                    yield f"{ws.title} r{i}c{j}", norm(v)
    wb.close()


def cells_docx(path):
    d = docx.Document(path)
    head = "(前言)"
    for k, p in enumerate(d.paragraphs, 1):
        style = str(p.style.name).lower() if p.style is not None else ""
        if style.startswith("heading") and p.text.strip():
            head = norm(p.text)
        if p.text.strip():
            yield f"para {k} [{head[:40]}]", norm(p.text)
    for ti, t in enumerate(d.tables, 1):
        for ri, row in enumerate(t.rows, 1):
            for ci, c in enumerate(row.cells, 1):
                if c.text.strip():
                    yield f"table {ti} r{ri}c{ci}", norm(c.text)


def sentence_of(text, idx):
    s = text.rfind(".", 0, idx) + 1
    e = text.find(".", idx)
    return text[s:e + 1 if e >= 0 else len(text)].strip()


def main():
    sources = [
        ("037", F037.name, cells_xlsx(F037)),
        ("SYS2", SYS2.name, cells_xlsx(SYS2, {"Basic Report"})),
        ("CFTS_020", CFTS.name, cells_docx(CFTS)),
        ("SYS3", SYS3.name, cells_docx(SYS3)),
    ]
    found = defaultdict(list)     # abbrev -> [(expansion, src, file, loc, quote)]
    for tag, fname, it in sources:
        for loc, text in it:
            for rx, order in ((PAREN_AFTER, "fwd"), (PAREN_BEFORE, "bwd")):
                for m in rx.finditer(text):
                    raw, ab = (m.group(1), m.group(2)) if order == "fwd" \
                        else (m.group(2), m.group(1))
                    words = norm(raw).split()
                    if order == "fwd":
                        exp, filler = longest_expansion(words, ab)
                    else:
                        ok, filler = initials_match(words, ab)
                        exp = " ".join(words) if ok else None
                    if not exp:
                        continue
                    if ab.upper() == exp.upper().replace(" ", ""):
                        continue          # "(HU)" restating itself
                    found[ab].append((exp, tag, fname, loc,
                                      sentence_of(text, m.start()),
                                      "filler-skipped" if filler else "strict"))

    print("# R-DM22 縮寫並列清點")
    print("兩種形態：`Expansion (ABBR)` 與 `ABBR (Expansion)`；逐字，"
          "不以領域常識建條目")
    print(f"掃描來源：037（全分頁全格）／SYS2 `Basic Report`（全格）／"
          f"CFTS_020（段落＋表格）／SYS3（段落＋表格）")
    print(f"共取得 {len(found)} 個縮寫之並列，"
          f"{sum(len(v) for v in found.values())} 處出現")

    # --- 停止條件 16：同一縮寫多種展開
    conflicts = {}
    for ab, hits in found.items():
        exps = {e.lower(): e for e, *_ in hits}
        if len(exps) > 1:
            conflicts[ab] = exps
    print(f"\n## 停止條件 16 檢查 —— 同一縮寫是否有兩種展開")
    if not conflicts:
        print("  無衝突。")
    for ab, exps in conflicts.items():
        print(f"  **{ab}**: {sorted(exps.values())}")
        for e, tag, fn, loc, q, mode in found[ab]:
            print(f"     [{tag} {loc}] {e!r} ({mode}) ← {q[:120]}")

    rows = []
    for ab in sorted(found):
        e, tag, fn, loc, q, mode = found[ab][0]
        usable = "Y" if len(e.split()) >= 2 else "N"
        rows.append((ab, e, fn, f"{tag} {loc}", q, usable,
                     str(len(found[ab])), mode))

    print("\n## 逐條（每個縮寫取其首處出現為出處，出現次數另計）")
    print("| abbrev | expansion | 首字母判準 | usable(>=2 詞) | 出現處數 "
          "| source_file | source_locator |")
    print("|---|---|---|---|---|---|---|")
    for ab, e, fn, loc, q, usable, cnt, mode in rows:
        print(f"| {ab} | {e} | {mode} | {usable} | {cnt} | {fn} | {loc} |")

    print("\n## 引句（cooccurrence_quote）")
    for ab, e, fn, loc, q, usable, cnt, mode in rows:
        print(f"\n### {ab} = {e}")
        print(f"  [{loc}] {q}")

    print("\n## 下放包 06 步驟 2 指名之縮寫 —— 逐一回報")
    for ab in MUST_CHECK:
        if ab in found:
            e = found[ab][0][0]
            print(f"  {ab}: 查得並列 → {e!r}"
                  f"（{len(found[ab])} 處，usable="
                  f"{'Y' if len(e.split()) >= 2 else 'N'}）")
        else:
            print(f"  {ab}: **查無並列** —— 依 R-DM22 不建條目")

    out = FEAT / "data" / "glossary.tsv"
    with out.open("w", encoding="utf-8") as fh:
        fh.write("abbrev\texpansion\tsource_file\tsource_locator\t"
                 "cooccurrence_quote\tusable\toccurrences\tinitials_rule\n")
        for r in rows:
            fh.write("\t".join(x.replace("\t", " ") for x in r) + "\n")
    write_meta(out, ["abbrev", "expansion", "source_file", "source_locator", "cooccurrence_quote", "usable", "occurrences", "initials_rule"], len(rows),
               generated_by="features/display/scripts/build_glossary.py",
               rulings=["R-DM22"],
               measurement_conditions="收錄判準為候選詞之首字母須逐一拼出該縮寫；initials_rule 欄記 strict／filler-skipped",
               notes="每條必引一處同句並列之來源；查無並列者不建條目。")
    print(f"\nwrote {out}")
    if conflicts:
        print("\n**停止條件 16 觸發** —— 見上方衝突清單，不擇一。")
        return 1
    return 0


if __name__ == "__main__":
    sys.exit(main())
