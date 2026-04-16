"""Tests for spec document parser module (RULES.md §14.5, §14.6)."""
import pytest
from openpyxl import Workbook
from docx import Document

from spec_parser import (
    detect_format,
    parse_pdf,
    parse_docx,
    parse_xlsx,
    merge_indexes,
)


@pytest.fixture
def sample_docx(tmp_path):
    """Create a sample DOCX with headings and PDM codes."""
    filepath = tmp_path / "spec.docx"
    doc = Document()
    doc.add_heading("Main Feature", level=1)
    doc.add_paragraph("PDM01 allows the user to access Device Manager.")
    doc.add_paragraph("PDM01.1 adds the shortcut to status bar.")
    doc.add_heading("Secondary Feature", level=2)
    doc.add_paragraph("TD1.7 handles touch display calibration.")
    doc.add_paragraph("No code in this paragraph.")
    doc.save(filepath)
    return str(filepath)


@pytest.fixture
def sample_xlsx_spec(tmp_path):
    """Create a sample XLSX spec file."""
    filepath = tmp_path / "spec.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.title = "Requirements"
    ws.cell(row=1, column=1, value="ID")
    ws.cell(row=1, column=2, value="Description")
    ws.cell(row=2, column=1, value="REQ-001")
    ws.cell(row=2, column=2, value="PDM05.2 power delivery setup")
    ws.cell(row=3, column=1, value="REQ-002")
    ws.cell(row=3, column=2, value="PDMS01 system configuration")
    ws.cell(row=4, column=1, value="REQ-003")
    ws.cell(row=4, column=2, value="No code here")
    wb.save(filepath)
    return str(filepath)


class TestDetectFormat:
    def test_pdf(self):
        assert detect_format("report.pdf") == "pdf"

    def test_docx(self):
        assert detect_format("spec.docx") == "docx"

    def test_xlsx(self):
        assert detect_format("data.xlsx") == "xlsx"

    def test_unknown(self):
        assert detect_format("file.txt") is None

    def test_case_insensitive(self):
        assert detect_format("FILE.PDF") == "pdf"


class TestParseDocx:
    def test_extracts_pdm_codes(self, sample_docx):
        index = parse_docx(sample_docx)
        assert "PDM01" in index
        assert "PDM01.1" in index
        assert "TD1.7" in index

    def test_includes_heading(self, sample_docx):
        index = parse_docx(sample_docx)
        assert index["PDM01"]["heading"] == "Main Feature"
        assert index["TD1.7"]["heading"] == "Secondary Feature"

    def test_includes_text(self, sample_docx):
        index = parse_docx(sample_docx)
        assert "access Device Manager" in index["PDM01"]["text"]

    def test_code_count(self, sample_docx):
        index = parse_docx(sample_docx)
        assert len(index) == 3


class TestParseXlsx:
    def test_extracts_pdm_codes(self, sample_xlsx_spec):
        index = parse_xlsx(sample_xlsx_spec)
        assert "PDM05.2" in index
        assert "PDMS01" in index

    def test_includes_row_info(self, sample_xlsx_spec):
        index = parse_xlsx(sample_xlsx_spec)
        assert index["PDM05.2"]["row"] == 2

    def test_includes_text(self, sample_xlsx_spec):
        index = parse_xlsx(sample_xlsx_spec)
        assert "power delivery" in index["PDM05.2"]["text"]

    def test_skips_no_code_rows(self, sample_xlsx_spec):
        index = parse_xlsx(sample_xlsx_spec)
        assert len(index) == 2


class TestMergeIndexes:
    def test_merge_both(self):
        sys1_index = {
            "PDM01": {
                "nrl_id": "NRL-100",
                "outline": "2.2",
                "source_id": "Source_2.2",
                "description": "PDM01 short desc",
            },
        }
        slot_c_index = {
            "PDM01": {
                "heading": "Main Feature",
                "text": "PDM01 full detailed text from spec document.",
            },
            "PDM02": {
                "heading": "Other",
                "text": "PDM02 only in slot C.",
            },
        }
        merged = merge_indexes(sys1_index, slot_c_index)
        # PDM01: has both source_id and full_text
        assert merged["PDM01"]["source_id"] == "Source_2.2"
        assert merged["PDM01"]["full_text"] == "PDM01 full detailed text from spec document."
        # PDM02: only slot C, no source_id
        assert merged["PDM02"]["source_id"] is None
        assert "PDM02 only in slot C" in merged["PDM02"]["full_text"]

    def test_sys1_only(self):
        sys1_index = {
            "PDM01": {
                "nrl_id": "NRL-100",
                "outline": "2.2",
                "source_id": "Source_2.2",
                "description": "short",
            },
        }
        merged = merge_indexes(sys1_index, None)
        assert merged["PDM01"]["source_id"] == "Source_2.2"
        assert merged["PDM01"]["full_text"] is None

    def test_slot_c_only(self):
        slot_c_index = {
            "PDM01": {"heading": "H", "text": "full text"},
        }
        merged = merge_indexes(None, slot_c_index)
        assert merged["PDM01"]["source_id"] is None
        assert merged["PDM01"]["full_text"] == "full text"

    def test_both_none(self):
        merged = merge_indexes(None, None)
        assert merged == {}
