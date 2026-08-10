"""Tests for the cross-document citation resolver in build_stla_map.py.

CFTS024 delegates behaviour to other documents in a scheme that looks
resolvable and is not: `{See CFTS019-718}` is a short id that appears in no
supplied file — not in CFTS019, and not even in CFTS024's own anchors, which
are 7-digit. Two things therefore have to hold, and both are easy to get
wrong:

- the citation has to be FOUND. The docx breaks the token with non-breaking
  spaces (`CFTS0\\xa019-718` is the literal run in §1.3.3), so the pattern that
  reads naturally misses exactly the citation that reaches a leaf.
- an unresolvable citation has to STAY unresolved. Ranked candidates are an
  aid to a ruling, never a resolution; a ruling that the files stop supporting
  has to fail loudly, the same discipline as the declared-id overrides.
"""
import importlib.util
import sys
from pathlib import Path

import docx
import pytest

ROOT = Path(__file__).resolve().parent.parent

_spec = importlib.util.spec_from_file_location(
    "amfm_stla_map_refs", ROOT / "AMFMHMI" / "scripts" / "build_stla_map.py")
if _spec is None or _spec.loader is None:
    pytest.skip("build_stla_map.py not present", allow_module_level=True)
bsm = importlib.util.module_from_spec(_spec)
sys.modules["amfm_stla_map_refs"] = bsm
_spec.loader.exec_module(bsm)


def clauses(*items) -> dict[int, dict]:
    return {cid: {"metadata": "", "text": text, "section": sec,
                  "section_title": title}
            for cid, text, sec, title in items}


TONE = clauses(
    (4866060, "The HU shall not activate another acceptance or rejection tone "
              "until the current location on the touchscreen is released.",
     "1.3.2.6", "Confirmation Tone Generation"),
    (4866062, "The key press rejection tone shall be applied for a short press "
              "event, if the function associated with the screen press is not "
              "allowed.", "1.3.2.6", "Confirmation Tone Generation"),
    (4866400, "The HU shall store the current mode settings when the user "
              "switches source.", "1.3.3.2", "Mode Settings"),
)

CITING = ("If all presets are deleted, HU shall play the rejection tone "
          "{See CFTS0\xa019-718} when user tries to access preset through the "
          "HU HMI or steering wheel controls.")


# ---------------------------------------------------------------- detection

def test_a_token_broken_by_non_breaking_spaces_is_still_found():
    """The §1.3.3 run reads `CFTS0\\xa019-718`; a plain pattern loses it."""
    hits = bsm.find_citations(CITING)
    assert [h["token"] for h in hits] == ["CFTS019-718"]


def test_the_citing_document_is_marked_when_it_cites_itself():
    hits = bsm.find_citations("see {See CFTS024-789} for HD Tuner", "CFTS024")
    assert hits[0]["self_reference"] is True
    assert hits[0]["doc"] == "CFTS024"


def test_a_clause_with_no_citation_yields_nothing():
    assert bsm.find_citations("HU shall display the browse category.") == []


def test_the_lead_phrase_stops_at_the_citation():
    """What follows the token describes the citing feature, not the cited one."""
    phrase = bsm.lead_phrase(CITING)
    assert phrase == "HU shall play the rejection tone"


# --------------------------------------------------------------- candidates

def test_candidates_rank_rare_shared_words_above_common_ones():
    """`rejection tone` decides; `the HU shall` is in every clause of a CFTS."""
    got = bsm._candidates(CITING, TONE, k=3)
    assert [c["clause_id"] for c in got][:2] == [4866060, 4866062]
    assert got[-1]["clause_id"] == 4866400


def test_the_citing_clause_is_not_offered_as_its_own_answer():
    """`behave as described in Section {See CFTS024-605}` ranks itself first."""
    self_citing = clauses(
        (4872507, "The Genre Seek function shall behave in the same manner as "
                  "the Seek Up function described in Section",
         "1.3.12", "Genre Seek"),
        (4872380, "The Seek Up function shall tune to the next receivable "
                  "station.", "1.3.1", "Seek Up"))
    got = bsm._candidates("shall behave as the Seek Up function {See CFTS024-605}",
                          self_citing, exclude={4872507})
    assert [c["clause_id"] for c in got] == [4872380]


# --------------------------------------------------------------- resolution

def paragraphs(*items):
    return {cid: {"metadata": "", "text": text} for cid, text in items}


MAPPING = {"SWE-RA-RAD-014": {"doc": "CFTS024", "stla_id": 4872420,
                              "section": "1.3.3"}}


def resolve(refs, overrides=None, paras=None):
    return bsm.resolve_citations(
        paras or paragraphs((4872420, CITING)), MAPPING, refs,
        overrides or {}, "CFTS024")


def test_an_unresolvable_token_stays_unresolved_with_candidates():
    got = resolve({"CFTS019": {"path": "cfts019.docx", "clauses": TONE}})
    entry = got["CFTS019-718"]
    assert entry["status"] == "unresolved-scheme-mismatch"
    assert "resolved_clause" not in entry
    assert entry["candidates"]


def test_the_leaf_that_owns_the_citing_clause_is_named():
    got = resolve({"CFTS019": {"path": "x", "clauses": TONE}})
    assert got["CFTS019-718"]["req_ids"] == ["SWE-RA-RAD-014"]


def test_an_absent_document_says_so_rather_than_guessing():
    got = resolve({})
    assert got["CFTS019-718"]["status"] == "document-not-supplied"
    assert "candidates" not in got["CFTS019-718"]


def test_a_token_that_is_a_real_anchor_resolves_without_a_ruling():
    paras = paragraphs((4872420, "see {See CFTS019-4866062} for the tone"))
    got = resolve({"CFTS019": {"path": "x", "clauses": TONE}}, paras=paras)
    entry = got["CFTS019-4866062"]
    assert entry["status"] == "resolved-by-anchor"
    assert entry["resolved_clause"] == 4866062


# ----------------------------------------------------------------- rulings

RULING = {"CFTS019-718": {"resolved_clause": 4866062, "ruling": "R11",
                          "evidence_phrase": "key press rejection tone"}}


def test_a_ruled_token_carries_the_clause_text_and_the_ruling():
    got = resolve({"CFTS019": {"path": "x", "clauses": TONE}}, RULING)
    entry = got["CFTS019-718"]
    assert entry["status"] == "resolved-by-ruling"
    assert entry["resolved_clause"] == 4866062
    assert entry["ruling"] == "R11"
    assert "key press rejection tone" in entry["resolved_text"]
    assert entry["resolved_section"] == "1.3.2.6"


def test_a_ruling_is_refused_when_the_clause_no_longer_carries_its_evidence():
    """CFTS019 reissued and 4866062 became something else."""
    moved = dict(TONE)
    moved[4866062] = dict(moved[4866062], text="The HU shall mute the speaker.")
    with pytest.raises(bsm.BuildError, match="does not hold"):
        resolve({"CFTS019": {"path": "x", "clauses": moved}}, RULING)


def test_a_ruling_onto_a_clause_that_does_not_exist_aborts():
    bad = {"CFTS019-718": dict(RULING["CFTS019-718"], resolved_clause=9999999)}
    with pytest.raises(bsm.BuildError, match="not a clause"):
        resolve({"CFTS019": {"path": "x", "clauses": TONE}}, bad)


def test_a_ruling_without_stated_evidence_aborts():
    bare = {"CFTS019-718": {"resolved_clause": 4866062, "ruling": "R11"}}
    with pytest.raises(bsm.BuildError, match="evidence_phrase"):
        resolve({"CFTS019": {"path": "x", "clauses": TONE}}, bare)


# -------------------------------------------------------- reference reading

def test_reference_clauses_carry_the_heading_they_sit_under(tmp_path):
    d = docx.Document()
    d.add_paragraph("1.3.2.6 Confirmation Tone Generation {4866050}",
                    style="Heading 3")
    d.add_paragraph("4866062: [Artifact Type:Subsystem Functional Requirement]")
    d.add_paragraph("The key press rejection tone shall be applied.")
    d.add_paragraph("A second paragraph of the same clause.")
    p = tmp_path / "cfts019.docx"
    d.save(p)

    got = bsm.load_reference_clauses(p)
    assert got[4866062]["section"] == "1.3.2.6"
    assert got[4866062]["section_title"] == "Confirmation Tone Generation"
    assert "second paragraph" in got[4866062]["text"]


def test_reference_anchors_need_not_increase(tmp_path):
    """Nothing brackets into a reference doc, so the ordering rule is not its rule."""
    d = docx.Document()
    d.add_paragraph("2 Later Chapter {4000000}", style="Heading 1")
    d.add_paragraph("4866062: [Artifact Type:SFR]")
    d.add_paragraph("text")
    d.add_paragraph("1 Earlier Chapter {5000000}", style="Heading 1")
    d.add_paragraph("4866063: [Artifact Type:SFR]")
    d.add_paragraph("text")
    p = tmp_path / "unordered.docx"
    d.save(p)

    got = bsm.load_reference_clauses(p)
    assert set(got) == {4866062, 4866063}
