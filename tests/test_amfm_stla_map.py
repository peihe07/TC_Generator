"""Tests for AMFMHMI/scripts/build_stla_map.py — the bracket mapping.

This map replaces hand-derivation, so the tests pin the two things a human
doing it by eye gets wrong, and the one thing a naive script gets wrong:

- full-width brackets: two of the 102 leaves tag their id as `（4942534）`
- out-of-range ids: 17 leaves belong to other CFTS documents, and a bracket
  lookup would silently pin every one of them to the last section of this one
- an unchecked batch table: the Phase 3 Test Set allocation is only worth
  something if a leaf landing outside its declared sections is an error
"""
import importlib.util
import sys
from pathlib import Path

import docx
import openpyxl
import pytest

ROOT = Path(__file__).resolve().parent.parent

_spec = importlib.util.spec_from_file_location(
    "amfm_stla_map", ROOT / "AMFMHMI" / "scripts" / "build_stla_map.py")
if _spec is None or _spec.loader is None:
    pytest.skip("build_stla_map.py not present", allow_module_level=True)
bsm = importlib.util.module_from_spec(_spec)
sys.modules["amfm_stla_map"] = bsm
_spec.loader.exec_module(bsm)


A03_HEADER = ["SWE-Requirement ID ", "Source Requirement ID",
              "Requirement  Title", "Requirement  Description",
              "Categorization"]


def a03(tmp_path, rows, name="a03.xlsx"):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Analysis Report"
    ws.append(["STLA Report_SWRA"])
    ws.append([])
    ws.append(A03_HEADER)
    for r in rows:
        ws.append(r)
    p = tmp_path / name
    wb.save(p)
    return p


def cfts(tmp_path, headings, name="spec.docx"):
    d = docx.Document()
    d.add_paragraph("front matter")
    for num, title, anchor in headings:
        d.add_paragraph(f"{num} {title} {{{anchor}}}", style="Heading 2")
        d.add_paragraph("body text for the section")
    p = tmp_path / name
    d.save(p)
    return p


SECTIONS = [("1.3", "HU Analog Tuner", 4872372),
            ("1.3.1", "Seek Up", 4872382),
            ("1.3.2", "Seek Down", 4872400)]


# ------------------------------------------------------------- id extraction

def test_full_width_brackets_are_read(tmp_path):
    """SWE-RA-RAD-087 and -097 tag their id with U+FF08/U+FF09."""
    p = a03(tmp_path, [
        ["SWE-RA-RAD-001", "s", "Ascii tagged (4872375)", "d", "Functional"],
        ["SWE-RA-RAD-002", "s", "Full width tagged（4872378）", "d", "Functional"],
    ])
    leaves = bsm.load_leaves(p)
    assert [l["stla_id"] for l in leaves] == [4872375, 4872378]


def test_a_leaf_with_no_id_aborts_rather_than_being_dropped(tmp_path):
    """A partial map is worse than none: the missing leaf gets no
    spec_reference and nothing says so."""
    p = a03(tmp_path, [
        ["SWE-RA-RAD-001", "s", "Tagged (4872375)", "d", "Functional"],
        ["SWE-RA-RAD-002", "s", "Untagged title", "d", "Functional"],
    ])
    with pytest.raises(bsm.BuildError, match="SWE-RA-RAD-002"):
        bsm.load_leaves(p)


def test_two_different_ids_in_one_title_abort(tmp_path):
    p = a03(tmp_path, [
        ["SWE-RA-RAD-001", "s", "Refers to (4872375) and also (4872999)",
         "d", "Functional"],
    ])
    with pytest.raises(bsm.BuildError, match="ambiguous"):
        bsm.load_leaves(p)


def test_the_same_id_repeated_is_not_ambiguous(tmp_path):
    p = a03(tmp_path, [
        ["SWE-RA-RAD-001", "s", "See (4872375) — restated (4872375)",
         "d", "Functional"],
    ])
    assert bsm.load_leaves(p)[0]["stla_id"] == 4872375


def test_non_functional_rows_are_not_leaves(tmp_path):
    p = a03(tmp_path, [
        ["SWE-RA-RAD-001", "s", "Tagged (4872375)", "d", "Functional"],
        ["SWE-RA-RAD-H1", "s", "A heading (4872376)", "d", "Heading"],
    ])
    assert [l["req_id"] for l in bsm.load_leaves(p)] == ["SWE-RA-RAD-001"]


# ---------------------------------------------------------------- sections

def test_headings_must_have_strictly_increasing_anchors(tmp_path):
    """Bracket lookup is undefined otherwise; the map would be quietly wrong."""
    p = cfts(tmp_path, [("1.1", "A", 4872400), ("1.2", "B", 4872372)])
    with pytest.raises(bsm.BuildError, match="strictly increasing"):
        bsm.load_sections(p)


def test_unanchored_headings_are_reported_not_used(tmp_path):
    d = docx.Document()
    d.add_paragraph("1.3 HU Analog Tuner {4872372}", style="Heading 2")
    d.add_paragraph("Appendix A — no anchor here", style="Heading 2")
    p = tmp_path / "s.docx"
    d.save(p)
    sections, unparsed = bsm.load_sections(p)
    assert [s["section"] for s in sections] == ["1.3"]
    assert any("Appendix" in u for u in unparsed)


# ----------------------------------------------------------------- bracket

def test_a_leaf_lands_in_the_largest_anchor_not_exceeding_it():
    secs = [{"section": n, "title": t, "anchor": a} for n, t, a in SECTIONS]
    assert bsm.bracket(secs, 4872382)["section"] == "1.3.1", "exact anchor hit"
    assert bsm.bracket(secs, 4872390)["section"] == "1.3.1", "between anchors"
    assert bsm.bracket(secs, 4872372)["section"] == "1.3"
    assert bsm.bracket(secs, 4872000) is None, "below the whole document"


def test_out_of_range_leaves_route_to_the_declared_external_doc():
    secs = [{"section": n, "title": t, "anchor": a} for n, t, a in SECTIONS]
    leaves = [{"req_id": "SWE-RA-RAD-001", "stla_id": 4872390, "title": "t",
               "description": "d", "source_components": "s"},
              {"req_id": "SWE-RA-RAD-097", "stla_id": 4939946, "title": "t",
               "description": "d", "source_components": "s"}]
    m = bsm.build(leaves, secs, {"CFTS004": ["SWE-RA-RAD-097"]}, "CFTS024")
    assert m["SWE-RA-RAD-001"]["doc"] == "CFTS024"
    assert m["SWE-RA-RAD-001"]["section"] == "1.3.1"
    assert m["SWE-RA-RAD-097"]["doc"] == "CFTS004"
    assert m["SWE-RA-RAD-097"]["resolution"] == "external-allocation"


def test_an_unallocated_out_of_range_leaf_aborts():
    """Without this the leaf silently pins to the document's last section."""
    secs = [{"section": n, "title": t, "anchor": a} for n, t, a in SECTIONS]
    leaves = [{"req_id": "SWE-RA-RAD-097", "stla_id": 4939946, "title": "t",
               "description": "d", "source_components": "s"}]
    with pytest.raises(bsm.BuildError, match="not\n?\\s*allocated|disagree"):
        bsm.build(leaves, secs, {}, "CFTS024")


def test_an_allocation_that_is_actually_in_range_aborts():
    """The ruled allocation and the measured range must agree both ways."""
    secs = [{"section": n, "title": t, "anchor": a} for n, t, a in SECTIONS]
    leaves = [{"req_id": "SWE-RA-RAD-001", "stla_id": 4872390, "title": "t",
               "description": "d", "source_components": "s"}]
    with pytest.raises(bsm.BuildError, match="disagree"):
        bsm.build(leaves, secs, {"CFTS011": ["SWE-RA-RAD-001"]}, "CFTS024")


# --------------------------------------------------------- id verification

CLAUSE_A = ("When tune down command is executed via the ICS hardcontrols "
            "(tune rotary knob), ICS shall send $ICS_KNOB2_DIR$ = [Decrement] "
            "and the HU shall perform the tune down function.")
CLAUSE_B = ("When the Cabin Output Channel has selected the HU Tuner source, "
            "the user can request to Tune to a Station by directly entering "
            "the Numeric Frequency of the Station.")


def _paras():
    return {4872451: {"metadata": "", "text": CLAUSE_A},
            4872457: {"metadata": "", "text": CLAUSE_B}}


def _leaf(rid, stla_id, title):
    return {rid: {"stla_id": stla_id, "title": title,
                  "spec_paragraph": _paras()[stla_id]["text"]}}


def test_a_leaf_pointing_at_the_wrong_clause_is_reported():
    """SWE-RA-RAD-029 declares 4872451 but describes 4872457's clause — a
    hand-typed id tail, invisible to the bracket map because it resolves
    perfectly to the wrong paragraph."""
    mapping = _leaf("SWE-RA-RAD-029", 4872451, CLAUSE_B + " (4872451)")
    found = bsm.verify_ids(mapping, _paras())
    assert len(found) == 1
    assert found[0]["declared_id"] == 4872451
    assert found[0]["better_id"] == 4872457
    assert found[0]["better_agreement"] > found[0]["declared_agreement"]


def test_a_leaf_pointing_at_its_own_clause_is_not_reported():
    mapping = _leaf("SWE-RA-RAD-028", 4872451, CLAUSE_A + " (4872451)")
    assert bsm.verify_ids(mapping, _paras()) == []


def test_a_heavily_paraphrased_leaf_is_not_called_a_wrong_id():
    """Low agreement with everything is paraphrase, not misattribution — the
    check needs a clearly better candidate, not just a poor declared one."""
    mapping = {"SWE-RA-RAD-050": {
        "stla_id": 4872451, "title": "Totally unrelated wording about presets",
        "spec_paragraph": CLAUSE_A}}
    assert bsm.verify_ids(mapping, _paras()) == []


def test_external_leaves_are_skipped():
    mapping = {"SWE-RA-RAD-097": {"stla_id": 4939946, "title": "t",
                                  "spec_paragraph": None}}
    assert bsm.verify_ids(mapping, _paras()) == []


# ----------------------------------------------------------- batch checking

BATCH_MD = """\
| Batch / Test Set | CFTS sections | n | Leaf ids | note |
|---|---|---|---|---|
| Seek | 1.3.1, 1.3.2 | 2 | 001-002 | x |
"""


def _mapping(section_for_002="1.3.2"):
    return {
        "SWE-RA-RAD-001": {"stla_id": 4872390, "section": "1.3.1",
                           "section_title": "Seek Up", "doc": "CFTS024"},
        "SWE-RA-RAD-002": {"stla_id": 4872410, "section": section_for_002,
                           "section_title": "x", "doc": "CFTS024"},
    }


def test_a_consistent_batch_table_reports_no_problems(tmp_path):
    md = tmp_path / "b.md"
    md.write_text(BATCH_MD, encoding="utf-8")
    assert bsm.check_batches(_mapping(), md) == []


def test_a_leaf_bracketing_outside_its_declared_sections_is_caught(tmp_path):
    md = tmp_path / "b.md"
    md.write_text(BATCH_MD, encoding="utf-8")
    problems = bsm.check_batches(_mapping(section_for_002="1.3.14"), md)
    assert any("SWE-RA-RAD-002" in p and "1.3.14" in p for p in problems)


def test_a_count_that_disagrees_with_the_id_list_is_caught(tmp_path):
    md = tmp_path / "b.md"
    md.write_text(BATCH_MD.replace("| 2 | 001-002", "| 3 | 001-002"),
                  encoding="utf-8")
    assert any("says 3 leaves" in p for p in bsm.check_batches(_mapping(), md))


def test_a_leaf_in_no_batch_is_caught(tmp_path):
    md = tmp_path / "b.md"
    md.write_text(BATCH_MD.replace("| 2 | 001-002", "| 1 | 001"), encoding="utf-8")
    assert any("in no batch" in p for p in bsm.check_batches(_mapping(), md))


def test_leaf_selector_expands_ranges_and_singletons():
    assert bsm.parse_leaf_selector("087, 089-091", "SWE-RA-RAD-") == [
        "SWE-RA-RAD-087", "SWE-RA-RAD-089", "SWE-RA-RAD-090", "SWE-RA-RAD-091"]
