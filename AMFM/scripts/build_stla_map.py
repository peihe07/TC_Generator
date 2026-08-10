#!/usr/bin/env python3
"""Step 1 (AMFM) — leaf → STLA id → CFTS section, by bracket mapping.

The 037's Categorization column is `Functional / NA` on all 102 leaves, so the
RD grouping axis carries no information (framework Part III, §4.1.2 degenerate
case). The axis that does carry it: every 037 Requirement Title ends with its
source STLA id, and every CFTS024 heading carries an `{id}` anchor that is
strictly increasing in document order. A leaf therefore lands in the section
whose anchor is the largest one not exceeding it — a bracket lookup, no
judgement.

Two traps this exists to stop being hand-work:

1. **Full-width brackets.** `SWE-RA-RAD-087` and `-097` tag their id as
   `（4942534）` (U+FF08/U+FF09), not `(4942534)`. An ASCII-only pattern loses
   exactly those two leaves and reports them as untagged.
2. **Out-of-range ids are not misses.** 17 leaves carry ids far above the
   CFTS024 anchor range because they belong to other CFTS documents. Bracket
   lookup would silently pin every one of them to the document's last section.
   Out-of-range is detected and routed by the ruled allocation instead
   (R7-Q3), and the two must agree on the same set of leaves or this fails.

Outputs (`data/`):
  stla_to_cfts.json  leaf -> {stla_id, doc, section, section_title, anchor}
  stla_to_cfts.tsv   the same, greppable, for review

Usage:
    python AMFM/scripts/build_stla_map.py --feature-dir AMFM
    python AMFM/scripts/build_stla_map.py --feature-dir AMFM --check-batches
"""
from __future__ import annotations

import argparse
import bisect
import difflib
import json
import re
import sys
from pathlib import Path

import docx
import openpyxl

sys.path.insert(0, str(Path(__file__).resolve().parent.parent.parent / "scripts"))
from feature_config import load_feature_config, resolve_path  # noqa: E402

# Both bracket families, because the 037 mixes them. Curly braces appear in
# the CFTS headings; the 037 uses round or full-width round.
STLA_ID_RE = re.compile(r"[({（]\s*(\d{7})\s*[)}）]")
HEADING_RE = re.compile(
    r"^(?P<num>\d+(?:\.\d+)*)\s+(?P<title>.+?)\s*[{(（]\s*(?P<anchor>\d{7})\s*[)}）]\s*$")
# Body requirements carry their own id as a leading tag on the metadata line;
# the requirement sentence is the paragraph(s) that follow it.
PARA_ANCHOR_RE = re.compile(r"^(\d{7})\s*[:：]\s*(.*)$")
A03_SHEET = "Analysis Report"
# Cross-document citations: `HU shall play the rejection tone {See CFTS019-718}`.
# The docx renders the token with non-breaking spaces inside it (`CFTS0\xa019-718`
# is the literal run in §1.3.3), so a plain `CFTS019` pattern misses exactly the
# citations that matter. Repair first, match second.
CITE_REPAIR_RE = re.compile(r"C\s*F\s*T\s*S\s*(\d)\s*(\d)\s*(\d)\s*[-–]\s*(\d{1,7})")
CITE_RE = re.compile(r"CFTS(\d{3})-(\d{1,7})")
HEADING_NUM_RE = re.compile(r"^(\d+(?:\.\d+)*)\s+(.+?)\s*$")


class BuildError(RuntimeError):
    pass


def norm(v) -> str:
    return re.sub(r"\s+", " ", str(v or "")).strip().lower()


# ------------------------------------------------------------------- inputs

def load_leaves(a03_path: Path) -> list[dict]:
    """037 leaves in document order: req_id, title, description, stla_id."""
    wb = openpyxl.load_workbook(a03_path, read_only=True)
    rows = list(wb[A03_SHEET].iter_rows(values_only=True))
    wb.close()
    hdr = next((i for i, r in enumerate(rows)
                if any("requirement description" in norm(v) for v in r)), None)
    if hdr is None:
        raise BuildError(f"{a03_path.name}: no header row")
    header = [norm(v) for v in rows[hdr]]

    def col(*need, forbid=()):
        hits = [i for i, v in enumerate(header)
                if all(t in v for t in need) and not any(t in v for t in forbid)]
        if len(hits) != 1:
            raise BuildError(f"{a03_path.name}: {need} matched {len(hits)} columns")
        return hits[0]

    ti, di = col("requirement", "title"), col("requirement", "description")
    ci = col("categorization", forbid=("sub",))
    si = col("source", "id")

    leaves, untagged, ambiguous = [], [], []
    for r in rows[hdr + 1:]:
        if not r[0] or not str(r[ci] or "").strip().lower().startswith("functional"):
            continue
        rid, title = str(r[0]).strip(), str(r[ti] or "")
        ids = STLA_ID_RE.findall(title)
        if not ids:
            untagged.append(rid)
            continue
        # The id is a trailing tag; a title quoting another id earlier would
        # be a genuine ambiguity, so distinct values are reported, not picked.
        if len(set(ids)) > 1:
            ambiguous.append((rid, sorted(set(ids))))
            continue
        leaves.append({"req_id": rid, "stla_id": int(ids[-1]),
                       "declared_stla_id": int(ids[-1]),
                       "title": title.strip(),
                       "description": str(r[di] or "").strip(),
                       "source_components": str(r[si] or "").strip()})
    if untagged or ambiguous:
        raise BuildError(
            "cannot read an STLA id from every leaf — this map must be total, "
            f"not best-effort. untagged={untagged} ambiguous={ambiguous}")
    return leaves


def load_paragraph_anchors(docx_path: Path) -> dict[int, dict]:
    """{stla id: {metadata, text}} for requirements written in the body.

    Every CFTS requirement paragraph is introduced by a metadata line tagged
    with its own id (`4872375: [Artifact Type:…] [State:…]`); the requirement
    sentence is what follows, up to the next tagged line or heading. This is a
    strictly finer resolution than the section bracket, and all 85 in-range
    leaves reach it — so a leaf's context can be its own clause rather than
    its chapter.
    """
    paras = docx.Document(docx_path).paragraphs
    tagged = [(i, PARA_ANCHOR_RE.match(p.text.strip()))
              for i, p in enumerate(paras)]
    tagged = [(i, m) for i, m in tagged if m]
    out = {}
    for k, (i, m) in enumerate(tagged):
        end = tagged[k + 1][0] if k + 1 < len(tagged) else len(paras)
        body = []
        for p in paras[i + 1:end]:
            if p.style.name.startswith("Heading"):
                break
            if p.text.strip():
                body.append(p.text.strip())
        out[int(m.group(1))] = {"metadata": m.group(2).strip(),
                                "text": "\n".join(body)}
    return out


def load_sections(docx_path: Path) -> tuple[list[dict], list[str]]:
    """CFTS headings in document order, with their anchors."""
    doc = docx.Document(docx_path)
    sections, unparsed = [], []
    for p in doc.paragraphs:
        if not p.style.name.startswith("Heading"):
            continue
        text = p.text.strip()
        if not text:
            continue
        m = HEADING_RE.match(text)
        if m:
            sections.append({"section": m.group("num"),
                             "title": m.group("title").strip(),
                             "anchor": int(m.group("anchor"))})
        else:
            unparsed.append(text)
    if not sections:
        raise BuildError(f"{docx_path.name}: no anchored headings found")
    anchors = [s["anchor"] for s in sections]
    breaks = [(sections[i]["section"], anchors[i], sections[i + 1]["section"],
               anchors[i + 1])
              for i in range(len(anchors) - 1) if anchors[i] >= anchors[i + 1]]
    if breaks:
        raise BuildError(
            "CFTS heading anchors are not strictly increasing in document "
            "order, so a bracket lookup is not defined. Breaks: "
            f"{breaks[:5]}")
    return sections, unparsed


# ------------------------------------------------------------------ mapping

def parse_leaf_selector(spec: str, prefix: str) -> list[str]:
    """`087, 089-096` -> full req_ids. Ranges are inclusive and numeric."""
    out = []
    for part in str(spec).split(","):
        part = part.strip()
        if not part:
            continue
        if "-" in part:
            a, b = (int(x) for x in part.split("-", 1))
            out.extend(f"{prefix}{n:03d}" for n in range(a, b + 1))
        else:
            out.append(f"{prefix}{int(part):03d}")
    return out


def bracket(sections: list[dict], stla_id: int) -> dict | None:
    """The section whose anchor is the largest not exceeding stla_id."""
    anchors = [s["anchor"] for s in sections]
    i = bisect.bisect_right(anchors, stla_id) - 1
    return sections[i] if i >= 0 else None


def build(leaves, sections, external: dict[str, list[str]], home_doc: str,
          paragraphs: dict[int, dict] | None = None,
          owning_docs: dict[str, dict] | None = None):
    """leaf -> its clause, in whichever document owns it.

    `owning_docs` supplies the clause index of each external document, so a
    leaf allocated to CFTS011 or CFTS004 resolves to a section and a clause the
    same way a CFTS024 leaf does. Without it those leaves stop at
    `external-allocation`: the doc is known, the clause is not, and the
    generator sees only the 037 title — which is how a whole batch ends up with
    unverifiable spec references.
    """
    lo, hi = sections[0]["anchor"], sections[-1]["anchor"]
    declared_external = {rid: doc for doc, ids in external.items() for rid in ids}

    mapping, out_of_range = {}, []
    for leaf in leaves:
        sid = leaf["stla_id"]
        inside = lo <= sid <= hi
        if not inside:
            out_of_range.append(leaf["req_id"])
        entry = {"stla_id": sid, "title": leaf["title"],
                 "description": leaf["description"],
                 "source_components": leaf["source_components"]}
        if inside:
            sec = bracket(sections, sid)
            para = (paragraphs or {}).get(sid)
            entry |= {"doc": home_doc, "section": sec["section"],
                      "section_title": sec["title"], "anchor": sec["anchor"],
                      "spec_paragraph": para["text"] if para else None,
                      "spec_paragraph_metadata": para["metadata"] if para else None,
                      "resolution": "paragraph" if para else "bracket"}
        else:
            doc = declared_external.get(leaf["req_id"])
            clause = ((owning_docs or {}).get(doc) or {}).get(sid)
            if clause:
                entry |= {"doc": doc, "section": clause["section"],
                          "section_title": clause["section_title"],
                          "anchor": None,
                          "spec_paragraph": clause["text"],
                          "spec_paragraph_metadata": clause["metadata"],
                          "resolution": "paragraph"}
            else:
                entry |= {
                    "doc": doc, "section": None, "section_title": None,
                    "anchor": None,
                    "resolution": ("external-allocation" if doc and
                                   doc not in (owning_docs or {})
                                   else "external-clause-not-found" if doc
                                   else "UNRESOLVED")}
        mapping[leaf["req_id"]] = entry

    # The mechanical test (is the id inside this document's anchor range?) and
    # the ruled allocation (which other document it belongs to) are derived
    # independently. They must select the same leaves; a disagreement means one
    # of them moved, and guessing which would silently mis-cite a spec.
    declared = set(declared_external)
    found = set(out_of_range)
    if declared != found:
        raise BuildError(
            "the ruled external allocation and the measured out-of-range set "
            f"disagree.\n  declared external, but inside {home_doc}: "
            f"{sorted(declared - found)}\n  outside {home_doc}, but not "
            f"allocated to any document: {sorted(found - declared)}")
    return mapping


def apply_overrides(leaves: list[dict], overrides: dict,
                    paragraphs: dict[int, dict]) -> list[dict]:
    """Apply ruled declared-id corrections, refusing ones that do not hold up.

    An override rewrites which clause a requirement traces to, so it is checked
    against the same evidence that justified it: the corrected clause must
    match the leaf's own text better than the declared one does. A ruling that
    stops being true of the files — the 037 reissued, the CFTS renumbered —
    then fails here instead of silently re-pointing a citation.
    """
    applied = []
    by_id = {l["req_id"]: l for l in leaves}
    for rid, spec in (overrides or {}).items():
        leaf = by_id.get(rid)
        if leaf is None:
            raise BuildError(f"stla_id_overrides names {rid}, which is not a "
                             "037 leaf")
        declared, corrected = int(spec["declared"]), int(spec["corrected"])
        if leaf["declared_stla_id"] != declared:
            raise BuildError(
                f"{rid}: override expects declared id {declared}, but the 037 "
                f"now says {leaf['declared_stla_id']} — the ruling "
                f"({spec.get('ruling', '?')}) predates this file")
        if corrected not in paragraphs:
            raise BuildError(f"{rid}: override target {corrected} is not a "
                             "clause in the primary document")
        was = difflib.SequenceMatcher(
            None, _fold(leaf["title"]),
            _fold(paragraphs.get(declared, {}).get("text", ""))).ratio()
        now = difflib.SequenceMatcher(
            None, _fold(leaf["title"]),
            _fold(paragraphs[corrected]["text"])).ratio()
        if now <= was:
            raise BuildError(
                f"{rid}: override to {corrected} agrees {now:.3f} with the "
                f"leaf, no better than the declared {declared} at {was:.3f} — "
                "the evidence for this ruling no longer holds")
        leaf["stla_id"] = corrected
        applied.append({"req_id": rid, "declared": declared,
                        "corrected": corrected, "ruling": spec.get("ruling"),
                        "agreement_declared": round(was, 3),
                        "agreement_corrected": round(now, 3)})
    return applied


def _fold(s: str) -> str:
    s = re.sub(r"[({（]\s*\d{7}\s*[)}）]", " ", str(s or ""))
    return re.sub(r"\s+", " ", re.sub(r"[^a-z0-9$ ]", " ", s.lower())).strip()


def verify_ids(mapping: dict, paragraphs: dict[str, dict],
               floor: float = 0.55, margin: float = 0.15) -> list[dict]:
    """Does each leaf's declared id point at the clause the leaf describes?

    `paragraphs` is {document: {clause id: clause}} — the check runs per
    document, so the 17 leaves owned by CFTS011 / CFTS004 are screened the same
    way the CFTS024 ones are.

    The id is a hand-typed tail on the 037 title, so it is copyable — and
    A-AM08 already records four suspected duplicates. A wrong id is invisible
    to the bracket map (it resolves perfectly, to the wrong clause) and ships
    as a wrong spec_reference, which is a traceability defect rather than a
    wording one.

    Reported when the declared paragraph is a poor match AND some other
    paragraph is clearly better. Both conditions matter: a leaf that
    paraphrases heavily scores low against everything and is not evidence of
    a wrong id.

    A token-overlap prefilter keeps this to a few seconds over ~1400
    paragraphs; difflib then ranks the survivors.
    """
    # One index per document: a leaf is only ever compared against the clauses
    # of the document that owns it. Comparing across documents would rank a
    # CFTS004 clause as the "better" match for a CFTS024 leaf, which is not a
    # finding anyone can act on.
    by_doc = {}
    for doc, clauses_ in paragraphs.items():
        by_doc[doc] = {pid: (_fold(p["text"]), set(_fold(p["text"]).split()))
                       for pid, p in clauses_.items()}
    findings = []
    for rid, e in mapping.items():
        index = by_doc.get(e.get("doc"), {})
        if not e.get("spec_paragraph") or e["stla_id"] not in index:
            continue
        want = _fold(e["title"])
        want_set = set(want.split())
        if not want_set:
            continue
        declared = difflib.SequenceMatcher(
            None, want, index[e["stla_id"]][0]).ratio()
        if declared >= floor:
            continue
        prefiltered = sorted(
            index, key=lambda pid: -len(want_set & index[pid][1]))[:20]
        best_pid, best = None, 0.0
        for pid in prefiltered:
            r = difflib.SequenceMatcher(None, want, index[pid][0]).ratio()
            if r > best:
                best_pid, best = pid, r
        if best_pid != e["stla_id"] and best - declared >= margin:
            findings.append({
                "req_id": rid, "declared_id": e["stla_id"],
                "declared_agreement": round(declared, 3),
                "better_id": best_pid, "better_agreement": round(best, 3),
            })
    return findings


def unallocated_clauses(mapping: dict, sections: list[dict],
                        paragraphs: dict[int, dict], home_doc: str) -> dict:
    """Per section: which clause ids no 037 leaf claims (A-AM10, RD-1 Q-AM3).

    R10-2 made absorption a ruled operation with a decision test — unallocated
    + same section + elaborates the cited clause → absorb and multi-cite;
    otherwise it is a coverage hole. Both halves need the same input: the list,
    per section, of what nobody claimed. Producing it by hand per batch is how
    the Tune batch nearly shipped without wrap-around.

    Only sections the 037 actually uses are reported. A section with no leaf at
    all is a different question (whole-capability omission) and is counted
    separately rather than mixed in.
    """
    anchors = [s["anchor"] for s in sections]
    by_num = {s["section"]: s for s in sections}

    def section_of(pid: int) -> str | None:
        i = bisect.bisect_right(anchors, pid) - 1
        return sections[i]["section"] if i >= 0 else None

    claimed = {e["stla_id"] for e in mapping.values() if e["doc"] == home_doc}
    # Section numbers are per document and they collide: CFTS011 and CFTS004
    # both restart at 1 and share 30 / 38 numbers with CFTS024. Filtering to the
    # primary document's leaves is what stops a CFTS011 section number from
    # marking the same-numbered CFTS024 section as "used" and dragging its
    # clauses into this report as unallocated.
    used = {e["section"] for e in mapping.values()
            if e["section"] and e["doc"] == home_doc}

    out: dict[str, dict] = {}
    for pid in sorted(paragraphs):
        sec = section_of(pid)
        if sec is None or sec not in used:
            continue
        meta, text = paragraphs[pid]["metadata"], paragraphs[pid]["text"]
        kind = ("SFR" if "Subsystem Functional Requirement" in meta
                else "Description" if "Description" in meta else "other")
        entry = out.setdefault(sec, {
            "section_title": by_num[sec]["title"],
            "claimed": [], "unallocated": []})
        if pid in claimed:
            entry["claimed"].append(pid)
        else:
            entry["unallocated"].append({
                "id": pid, "kind": kind,
                # The scoping tags decide whether a clause is even this ECU's
                # behaviour; carrying them means the reader classifies without
                # reopening the docx.
                "scope": re.sub(r"\s+", " ", meta)[:200],
                "text": re.sub(r"\s+", " ", text)[:400],
            })
    return out


# ------------------------------------------------- cross-document citations

def repair_citations(text: str) -> str:
    """Rejoin citation tokens the docx broke apart with non-breaking spaces."""
    return CITE_REPAIR_RE.sub(r"CFTS\1\2\3-\4", str(text or "").replace("\xa0", " "))


def find_citations(text: str, home_doc: str | None = None) -> list[dict]:
    """`{See CFTS019-718}` occurrences, with the wording that introduces them.

    The surrounding sentence is carried because the cited id is in the citing
    document's own cross-reference scheme, not the STLA anchor scheme (see
    `resolve_citations`) — when the token cannot be resolved by id, the phrase
    it hangs off ("HU shall play the rejection tone") is the only evidence
    left for finding the clause it means.
    """
    fixed = repair_citations(text)
    out = []
    for m in CITE_RE.finditer(fixed):
        doc = f"CFTS{m.group(1)}"
        out.append({
            "token": f"{doc}-{m.group(2)}", "doc": doc, "cited_id": m.group(2),
            "self_reference": doc == home_doc,
            "context": re.sub(r"\s+", " ",
                              fixed[max(0, m.start() - 160):m.end() + 60]).strip(),
        })
    return out


def load_reference_clauses(docx_path: Path) -> dict[int, dict]:
    """A cited document's clauses, each tagged with the heading it sits under.

    Same anchor convention as the primary document, but read in one pass with
    the section carried along: a reference doc is consulted by clause, and a
    bare clause id with no chapter is not something a reviewer can check.
    Unlike `load_sections` this does not require anchors to increase — nothing
    brackets into a reference document, so the ordering invariant does not apply.
    """
    paras = docx.Document(docx_path).paragraphs
    out: dict[int, dict] = {}
    section = ("", "")
    current: dict | None = None
    for p in paras:
        text = p.text.strip()
        if not text:
            continue
        if p.style.name.startswith("Heading"):
            head = HEADING_RE.match(text) or HEADING_NUM_RE.match(text)
            if head:
                num, title = head.group(1), head.group(2)
                section = (num, title.strip())
            current = None
            continue
        m = PARA_ANCHOR_RE.match(text)
        if m:
            current = {"metadata": m.group(2).strip(), "text": "",
                       "section": section[0], "section_title": section[1]}
            out[int(m.group(1))] = current
        elif current is not None:
            current["text"] = (current["text"] + "\n" + text).strip()
    return out


def lead_phrase(context: str) -> str:
    """The wording the citation hangs off: `HU shall play the rejection tone`.

    Everything after the token describes the citing feature, not the cited
    behaviour ("…when user tries to access preset through the HU HMI"), and
    dragging it into the match pulls every candidate toward radio vocabulary.
    """
    before = re.split(r"\{\s*See|CFTS\d{3}-", context)[0]
    tail = re.split(r"[.;,]", before)[-1].strip()
    return tail if len(tail) >= 15 else before.strip()


def _candidates(context: str, clauses: dict[int, dict], k: int = 3,
                exclude: set[int] | None = None) -> list[dict]:
    """Clauses in the cited document that best match the citing phrase.

    Ranked by shared words weighted by how rare each is in the cited document,
    not by raw overlap: `the HU shall` is in most clauses of every CFTS and
    scores four unrelated candidates identically, while `rejection` occurs in
    three. difflib then breaks ties on wording.
    """
    want = set(_fold(lead_phrase(context)).split())
    if not want:
        return []
    # A clause citing a document is never the answer to its own citation; left
    # in, `see Section {CFTS024-605}` ranks the sentence that wrote it first.
    index = {cid: _fold(c["text"]) for cid, c in clauses.items()
             if c["text"] and cid not in (exclude or set())}
    df: dict[str, int] = {}
    for text in index.values():
        for w in set(text.split()):
            df[w] = df.get(w, 0) + 1
    n = max(len(index), 1)

    def score(cid: str) -> tuple[float, float]:
        shared = want & set(index[cid].split())
        weight = sum(1.0 / df.get(w, n) for w in shared)
        return (round(weight, 4),
                difflib.SequenceMatcher(None, " ".join(sorted(want)),
                                        index[cid]).ratio())

    ranked = sorted(index, key=score, reverse=True)[:k]
    return [{"clause_id": cid, "match_weight": score(cid)[0],
             "matched_terms": sorted(want & set(index[cid].split())),
             "section": clauses[cid]["section"],
             "section_title": clauses[cid]["section_title"],
             "text": re.sub(r"\s+", " ", clauses[cid]["text"])[:300]}
            for cid in ranked]


def _apply_citation_ruling(token: str, entry: dict, spec: dict,
                           clauses: dict[int, dict]):
    """A ruled token -> clause mapping, refused if the file stopped supporting it."""
    target = int(spec["resolved_clause"])
    if target not in clauses:
        raise BuildError(
            f"clause_citation_overrides[{token}] points at {target}, "
            f"which is not a clause in {entry['doc']}")
    phrase = _fold(spec.get("evidence_phrase", ""))
    if not phrase:
        raise BuildError(
            f"clause_citation_overrides[{token}] needs an "
            "evidence_phrase — the wording that makes this the right clause")
    if phrase not in _fold(clauses[target]["text"]):
        raise BuildError(
            f"clause_citation_overrides[{token}]: clause {target} no "
            f"longer contains {spec['evidence_phrase']!r} — the evidence for "
            f"this ruling ({spec.get('ruling', '?')}) does not hold")
    entry |= {"status": "resolved-by-ruling", "resolved_clause": target,
              "ruling": spec.get("ruling"),
              "resolved_section": clauses[target]["section"],
              "resolved_section_title": clauses[target]["section_title"],
              "resolved_text": re.sub(r"\s+", " ", clauses[target]["text"])}


def resolve_citations(clause_sources: dict[str, dict], mapping: dict,
                      references: dict[str, dict], overrides: dict,
                      home_doc: str) -> dict:
    """Every cross-document citation in the spec, and what it does or does not resolve to.

    The citing scheme is NOT the STLA anchor scheme: CFTS024 writes
    `{See CFTS019-718}` and `{See CFTS024-789}` — short ids that appear in no
    supplied document, including CFTS024's own. So a citation is quoted, never
    renumbered: the token is the reference, exactly as the source writes it.
    Where the cited document IS supplied, the candidate clauses are ranked so
    the mapping becomes a one-off ruling per token rather than a judgement call
    per test case; until that ruling exists the status says so, and the
    generator is told not to invent the referenced behaviour.
    """
    # Keyed by (document, clause) because clause ids are only unique within a
    # document, and the sweep now covers every document that owns leaves — a
    # CFTS004 clause citing `CFTS004-1316` is the same R11 case as a CFTS024
    # one, and scanning the primary alone would hide it from the batch that
    # actually tests that leaf.
    leaves_by_clause: dict[tuple[str, int], list[str]] = {}
    for rid, e in mapping.items():
        leaves_by_clause.setdefault((e["doc"], e["stla_id"]), []).append(rid)

    out: dict[str, dict] = {}
    for citing_doc, clauses_ in sorted(clause_sources.items()):
        for cid in sorted(clauses_):
            for hit in find_citations(clauses_[cid]["text"], citing_doc):
                entry = out.setdefault(hit["token"], {
                    "doc": hit["doc"], "cited_id": hit["cited_id"],
                    "status": "pending", "citing_clauses": [], "req_ids": []})
                entry["citing_clauses"].append({"citing_doc": citing_doc,
                                                "clause_id": cid,
                                                "context": hit["context"]})
                entry["req_ids"] = sorted(
                    set(entry["req_ids"])
                    | set(leaves_by_clause.get((citing_doc, cid), [])))

    for token, entry in out.items():
        clauses = (references.get(entry["doc"]) or {}).get("clauses")
        spec = (overrides or {}).get(token)
        if clauses is None:
            entry["status"] = "document-not-supplied"
        elif int(entry["cited_id"]) in clauses:
            entry |= {"status": "resolved-by-anchor",
                      "resolved_clause": int(entry["cited_id"])}
        elif spec:
            _apply_citation_ruling(token, entry, spec, clauses)
        else:
            entry["status"] = "unresolved-scheme-mismatch"
            entry["candidates"] = _candidates(
                entry["citing_clauses"][0]["context"], clauses,
                exclude={c["clause_id"] for c in entry["citing_clauses"]})
    return out


def check_batches(mapping: dict, batches_md: Path) -> list[str]:
    """Every leaf in a batch must bracket into one of that batch's sections.

    The Test Set table was derived by hand at Phase 3. This turns it into a
    checked claim: the batch doc says which CFTS sections a set covers, the
    bracket map says which section each leaf actually lands in.
    """
    row = re.compile(r"^\|\s*([^|]+?)\s*\|\s*([^|]*?)\s*\|\s*(\d+)\s*\|"
                     r"\s*([^|]+?)\s*\|")
    problems, seen = [], set()
    for line in batches_md.read_text(encoding="utf-8").splitlines():
        m = row.match(line)
        if not m or m.group(3) == "n":
            continue
        name, secs, count, ids = m.groups()
        req_ids = parse_leaf_selector(ids, "SWE-RA-RAD-")
        if len(req_ids) != int(count):
            problems.append(f"{name}: table says {count} leaves, lists {len(req_ids)}")
        declared = {s.strip() for s in re.split(r"[,;]", secs)
                    if re.match(r"^\s*\d+(\.\d+)*\s*$", s)}
        for rid in req_ids:
            seen.add(rid)
            entry = mapping.get(rid)
            if entry is None:
                problems.append(f"{name}: {rid} is not a 037 leaf")
            elif entry["section"] and declared and entry["section"] not in declared:
                problems.append(
                    f"{name}: {rid} (id {entry['stla_id']}) brackets into "
                    f"{entry['section']} '{entry['section_title']}', which the "
                    f"table does not list ({sorted(declared)})")
    missing = sorted(set(mapping) - seen)
    if missing:
        problems.append(f"{len(missing)} leaves are in no batch: {missing[:10]}")
    return problems


# -------------------------------------------------------------------- main

def run(args) -> int:
    cfg = load_feature_config(args.feature_dir)
    root = Path(args.feature_dir)
    a03 = resolve_path(cfg, "a03_report")
    cfts = resolve_path(cfg, "cfts_doc")

    spec_cfg = cfg.get("spec_docs") or {}
    home_doc = spec_cfg.get("primary")
    external = {d: parse_leaf_selector(v, "SWE-RA-RAD-")
                for d, v in (spec_cfg.get("external") or {}).items()}
    if not home_doc:
        raise BuildError("feature.yaml needs spec_docs.primary (e.g. CFTS024)")

    # Every non-primary document, read once and used for both roles: resolving
    # the leaves an external document owns, and ranking candidates for the
    # clauses that cite it.
    other_docs = {}
    for doc, key in (spec_cfg.get("docs") or {}).items():
        if not cfg["paths"].get(key):
            other_docs[doc] = {"path": None, "clauses": None}
            continue
        path = resolve_path(cfg, key)
        other_docs[doc] = {"path": path.name,
                           "clauses": load_reference_clauses(path)}
    owning = {d: v["clauses"] for d, v in other_docs.items()
              if d in external and v["clauses"]}

    leaves = load_leaves(a03)
    sections, unparsed = load_sections(cfts)
    paragraphs = load_paragraph_anchors(cfts)
    overrides = apply_overrides(leaves, cfg.get("stla_id_overrides"), paragraphs)
    mapping = build(leaves, sections, external, home_doc, paragraphs, owning)
    for o in overrides:
        mapping[o["req_id"]] |= {
            "declared_stla_id": o["declared"], "id_override": o}

    data = root / "data"
    data.mkdir(exist_ok=True)
    (data / "stla_to_cfts.json").write_text(
        json.dumps(mapping, ensure_ascii=False, indent=1), encoding="utf-8")
    lines = ["req_id\tstla_id\tdoc\tsection\tsection_title\tresolution"]
    for rid, e in mapping.items():
        lines.append(f"{rid}\t{e['stla_id']}\t{e['doc']}\t{e['section'] or ''}"
                     f"\t{e['section_title'] or ''}\t{e['resolution']}")
    (data / "stla_to_cfts.tsv").write_text("\n".join(lines) + "\n", encoding="utf-8")

    by_doc: dict[str, int] = {}
    for e in mapping.values():
        by_doc[e["doc"]] = by_doc.get(e["doc"], 0) + 1
    print(f"leaves        : {len(mapping)}")
    print(f"{home_doc} sections: {len(sections)} anchored headings "
          f"({sections[0]['anchor']}–{sections[-1]['anchor']})")
    if unparsed:
        print(f"headings without an anchor (not usable as brackets): "
              f"{len(unparsed)}")
    for doc, n in sorted(by_doc.items(), key=lambda kv: -kv[1]):
        print(f"  {doc}: {n} leaves")
    used_by_doc: dict[str, set] = {}
    for e in mapping.values():
        if e["section"]:
            used_by_doc.setdefault(e["doc"], set()).add(e["section"])
    print("distinct sections used: " + ", ".join(
        f"{d}={len(s)}" for d, s in sorted(used_by_doc.items())))
    by_res: dict[str, int] = {}
    for e in mapping.values():
        by_res[e["resolution"]] = by_res.get(e["resolution"], 0) + 1
    print(f"paragraph anchors in {home_doc}: {len(paragraphs)}")
    print("resolution    : " + ", ".join(f"{k}={v}" for k, v in sorted(by_res.items())))
    coarse = [rid for rid, e in mapping.items() if e["resolution"] == "bracket"]
    if coarse:
        print(f"  section-level only (no paragraph anchor for the id): {coarse}")

    for o in overrides:
        print(f"id override   : {o['req_id']} {o['declared']} -> "
              f"{o['corrected']} ({o['ruling']}); agreement "
              f"{o['agreement_declared']} -> {o['agreement_corrected']}")
    unalloc = unallocated_clauses(mapping, sections, paragraphs, home_doc)
    (data / "unallocated_clauses.json").write_text(
        json.dumps(unalloc, ensure_ascii=False, indent=1), encoding="utf-8")
    n_un = sum(len(v["unallocated"]) for v in unalloc.values())
    n_sfr = sum(1 for v in unalloc.values() for c in v["unallocated"]
                if c["kind"] == "SFR")
    print(f"unallocated   : {n_un} clauses ({n_sfr} SFR) across "
          f"{len(unalloc)} used sections -> data/unallocated_clauses.json")
    worst = sorted(unalloc.items(), key=lambda kv: -len(kv[1]["unallocated"]))[:5]
    for sec, v in worst:
        print(f"  §{sec:14} {len(v['unallocated']):3} unallocated / "
              f"{len(v['claimed'])} claimed  {v['section_title'][:40]}")

    # Documents the spec cites but that own no leaf (CFTS019 Audio Management,
    # CFTS028 Voice Recognition). Declared in feature.yaml so an absent file is
    # a stated gap rather than a silent one.
    # The primary document cites itself in the same short-id scheme
    # (`{See CFTS024-789}`), so it is a citable source like any other — without
    # this those tokens would report as "document not supplied" while the file
    # sits in inputs/.
    references = dict(other_docs)
    references[home_doc] = {"path": cfts.name,
                            "clauses": load_reference_clauses(cfts)}
    citation_sources = {home_doc: paragraphs} | owning
    citations = resolve_citations(citation_sources, mapping, references,
                                  cfg.get("clause_citation_overrides"), home_doc)
    (data / "cross_doc_citations.json").write_text(
        json.dumps({t: {k: v for k, v in e.items()}
                    for t, e in sorted(citations.items())},
                   ensure_ascii=False, indent=1), encoding="utf-8")
    in_scope = {t: e for t, e in citations.items() if e["req_ids"]}
    by_status: dict[str, int] = {}
    for e in citations.values():
        by_status[e["status"]] = by_status.get(e["status"], 0) + 1
    print(f"cross-doc refs: {len(citations)} cited tokens "
          f"({len(in_scope)} reached by a leaf) -> data/cross_doc_citations.json")
    print("  status      : " + ", ".join(f"{k}={v}" for k, v in sorted(by_status.items())))
    for token, e in sorted(in_scope.items()):
        head = (f"-> {e['resolved_clause']}" if e.get("resolved_clause")
                else f"[{e['status']}]")
        print(f"  {token:16} {head:26} {', '.join(e['req_ids'])}")

    suspect = verify_ids(mapping, citation_sources)
    if suspect:
        print("\nDECLARED STLA ID DOES NOT MATCH THE CLAUSE (A-AM08 class):",
              file=sys.stderr)
        for f in suspect:
            print(f"  - {f['req_id']}: declared {f['declared_id']} "
                  f"(agreement {f['declared_agreement']}), but {f['better_id']} "
                  f"matches at {f['better_agreement']}", file=sys.stderr)
        print("  These are reported, never corrected here — R7-Q4 leaves each "
              "pair to Pei.", file=sys.stderr)
        (data / "stla_id_suspects.json").write_text(
            json.dumps(suspect, ensure_ascii=False, indent=1), encoding="utf-8")

    if args.check_batches:
        md = root / args.batches_md
        if not md.exists():
            raise BuildError(f"missing {md}")
        problems = check_batches(mapping, md)
        if problems:
            print("\nBATCH TABLE DISAGREES WITH THE BRACKET MAP:", file=sys.stderr)
            for p in problems:
                print(f"  - {p}", file=sys.stderr)
            return 1
        print(f"batch table   : consistent with the bracket map ({md.name})")
    return 0


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--feature-dir", default=".")
    ap.add_argument("--batches-md", default="docs/batches-amfm.md")
    ap.add_argument("--check-batches", action="store_true")
    args = ap.parse_args()
    try:
        return run(args)
    except BuildError as exc:
        print(f"ABORTED: {exc}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    sys.exit(main())
